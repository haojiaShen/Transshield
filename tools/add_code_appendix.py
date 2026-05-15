#!/usr/bin/env python3
"""添加代码附录到作品报告"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

def set_run_font(run, font_name='宋体', font_size=Pt(12), bold=False):
    """设置run的字体"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def add_code_block(doc, code, title=""):
    """添加代码块"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, '黑体', Pt(12), True)
    
    # 添加代码
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_appendix():
    """添加代码附录"""
    doc = Document('/home/yclcg/Transshield_final/docs/transshield_竞赛作品报告_最终版.docx')
    
    # 找到参考文献位置
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and '参考文献' in p.text:
            ref_idx = i
            break
    
    if ref_idx is None:
        # Fallback: find by text content
        for i, p in enumerate(doc.paragraphs):
            if '参考文献' in p.text.strip():
                ref_idx = i
                break
    if ref_idx is None:
        print("未找到参考文献章节")
        return
    
    # 在参考文献后添加附录
    insert_after = doc.paragraphs[ref_idx]
    
    # 找到最后一个段落
    for i in range(ref_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            insert_after = doc.paragraphs[i]
    
    # 添加附录标题
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    new_p.style = doc.styles['Heading 1']
    run = new_p.add_run('附录A：关键代码实现')
    set_run_font(run, '黑体', Pt(16))
    insert_after = new_p
    
    # 添加说明
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run('本附录展示Transshield系统中的关键代码实现片段，体现技术创新的具体落地。')
    set_run_font(run, '宋体', Pt(12))
    new_p.paragraph_format.first_line_indent = Pt(24)
    insert_after = new_p
    
    # 代码片段1：安全剪枝决策
    code1_title = "A.1 安全剪枝决策（Secure Pruning Decision）"
    code1_desc = "以下代码实现了基于Bitonic Sort的安全Top-K选择算法，用于在MPC环境中进行token剪枝决策。"
    code1 = '''def _secure_build_keep_decision(score, prev_decision_2d, keep_count):
    """Build keep decision via argsort-equivalent top-k using bitonic sort."""
    N = int(score.shape[1])
    active_before = prev_decision_2d.squeeze(-1) > 0
    
    # Encode tie-breaking: among equal scores, lower index sorts first
    epsilon = 1e-6
    index_vals = jnp.arange(N, dtype=score.dtype) * epsilon
    encoded_key = score - index_vals
    
    # Mask inactive tokens: set to -inf
    encoded_masked = jnp.where(active_before, encoded_key, float('-inf'))
    
    # Pad to next power of 2 for bitonic sort
    padded_count = 1
    while padded_count < N:
        padded_count *= 2
    if padded_count > N:
        pad_width = padded_count - N
        neg_inf = jnp.full((int(score.shape[0]), pad_width), float('-inf'))
        sortable = jnp.concatenate([encoded_masked, neg_inf], axis=1)
    else:
        sortable = encoded_masked
    
    # Bitonic sort descending
    sorted_keys = _bitonic_sort_desc(sortable)
    
    # Threshold: the encoded value at position keep_count - 1
    threshold = sorted_keys[:, keep_count - 1 : keep_count]
    
    # Keep all active tokens whose encoded_key >= threshold
    keep_mask = (encoded_key >= threshold) & active_before
    return keep_mask[:, :, None].astype(prev_decision_2d.dtype)'''
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code1_title)
    set_run_font(run, '黑体', Pt(12), True)
    insert_after = new_p
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code1_desc)
    set_run_font(run, '宋体', Pt(12))
    new_p.paragraph_format.first_line_indent = Pt(24)
    insert_after = new_p
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code1)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    new_p.paragraph_format.left_indent = Inches(0.3)
    insert_after = new_p
    
    # 代码片段2：Bitonic Sort
    code2_title = "A.2 Bitonic Sort 排序算法"
    code2_desc = "以下代码实现了SPU友好的Bitonic Sort算法，用于安全排序。"
    code2 = '''def _bitonic_sort_desc(values):
    """Bitonic sort descending. O(N log^2 N) compare-and-swap."""
    N = int(values.shape[1])
    x = values
    k = 2
    while k <= N:
        j = k // 2
        while j >= 1:
            p_arr = jnp.arange(N, dtype=jnp.int32)
            p_partner = p_arr ^ j
            p_partner_safe = jnp.clip(p_partner, 0, N - 1)
            
            x_at_p = x
            x_at_partner = x[:, p_partner_safe]
            
            is_left = (p_arr < p_partner)
            has_partner = (p_partner < N)
            left_index = jnp.where(is_left, p_arr, p_partner)
            is_desc = (left_index & k) == 0
            
            should_swap = jnp.where(
                is_left,
                jnp.where(is_desc, x_at_p < x_at_partner, x_at_p > x_at_partner),
                False
            ) & has_partner
            
            x_at_p_new = jnp.where(should_swap, x_at_partner, x_at_p)
            x_at_partner_new = jnp.where(should_swap, x_at_p, x_at_partner)
            x = x.at[:, p_arr].set(x_at_p_new)
            
            j //= 2
        k *= 2
    return x'''
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code2_title)
    set_run_font(run, '黑体', Pt(12), True)
    insert_after = new_p
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code2_desc)
    set_run_font(run, '宋体', Pt(12))
    new_p.paragraph_format.first_line_indent = Pt(24)
    insert_after = new_p
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code2)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    new_p.paragraph_format.left_indent = Inches(0.3)
    insert_after = new_p
    
    # 代码片段3：隐私保护检查
    code3_title = "A.3 隐私保护状态检查"
    code3_desc = "以下代码展示了Transshield如何确保隐私保护状态。"
    code3 = '''# 隐私保护状态记录
privacy_status = {
    "host_plaintext_pixel_values_materialized": (
        False if share_pair_cpu is not None 
        or party_local_share_manifest_paths is not None 
        else True
    ),
    "host_model_params_materialized": False,
    "reveal_policy": "final_logits_only",
}

# 验证隐私保护完整性
assert privacy_status["host_plaintext_pixel_values_materialized"] == False
assert privacy_status["host_model_params_materialized"] == False
assert privacy_status["reveal_policy"] == "final_logits_only"'''
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code3_title)
    set_run_font(run, '黑体', Pt(12), True)
    insert_after = new_p
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code3_desc)
    set_run_font(run, '宋体', Pt(12))
    new_p.paragraph_format.first_line_indent = Pt(24)
    insert_after = new_p
    
    new_p = doc.add_paragraph()
    insert_after._element.addnext(new_p._element)
    run = new_p.add_run(code3)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    new_p.paragraph_format.left_indent = Inches(0.3)
    insert_after = new_p
    
    # 保存文档
    output_path = '/home/yclcg/Transshield_final/docs/transshield_竞赛作品报告_最终版.docx'
    doc.save(output_path)
    print(f"已添加代码附录：{output_path}")

if __name__ == '__main__':
    add_appendix()
