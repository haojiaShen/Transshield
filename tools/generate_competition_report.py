#!/usr/bin/env python3
"""Generate the second-revision competition report aligned to current report evidence."""

from __future__ import annotations

import json
import shutil
import textwrap
from datetime import datetime
from pathlib import Path
from hashlib import sha256

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs" / "transshield_竞赛作品报告_第二次修订版.docx"
ASSET_DIR = REPO_ROOT / "docs" / "report_evidence" / "assets"
RESULT_DIR = REPO_ROOT / "results" / "report_evidence"
BEST_DEMO_JSON = REPO_ROOT / "artifacts" / "web_demo_assets" / "demo_content_summary.json"
COMM_PROFILE_JSON = RESULT_DIR / "mainline_communication_profile.json"
MEDICAL_DYNAMIC_SUMMARY_JSON = (
    RESULT_DIR
    / "medical_dynamic_threshold_calibration"
    / "dynamic_fullval_depth10_threshold_summary.json"
)
PRIMARY_PROTOCOL_FUZZ_JSON = RESULT_DIR / "protocol_fuzz_evidence.json"
PRIMARY_GUARD_STRESS_JSON = RESULT_DIR / "control_plane_guard_evidence.json"
CHANGE_AUDIT_DOC = REPO_ROOT / "docs" / "report_evidence" / "10_web_demo_change_audit.md"
MODIFICATION_AUDIT_DOC = REPO_ROOT / "docs" / "transshield_bumblebee_spu_modifications.md"
SPU_LICENSE = REPO_ROOT / "spu_vendored" / "LICENSE"
SPU_MODIFICATIONS = REPO_ROOT / "spu_vendored" / "MODIFICATIONS.md"
THIRD_PARTY = REPO_ROOT / "THIRD_PARTY.md"
REPRODUCE_DOC = REPO_ROOT / "README_REPRODUCE.md"
VENDORED_NOTICE_FILES = [
    REPO_ROOT / "spu_vendored" / "libspu" / "spu.proto",
    REPO_ROOT / "spu_vendored" / "libspu" / "mpc" / "cheetah" / "arithmetic.h",
    REPO_ROOT / "spu_vendored" / "libspu" / "mpc" / "cheetah" / "arithmetic.cc",
    REPO_ROOT / "spu_vendored" / "libspu" / "mpc" / "cheetah" / "protocol.cc",
]

CALIBRATION_FIGURE = ASSET_DIR / "medical_threshold_calibration_shift.png"
TOPOLOGY_FIGURE = ASSET_DIR / "system_trust_boundary_topology.png"
CURRENT_ARCH_FIGURE = REPO_ROOT / "2-1.png"
SEQUENCE_FIGURE = ASSET_DIR / "software_flow_sequence.png"
ROBUSTNESS_FIGURE = ASSET_DIR / "robustness_guard_matrix.png"
REPORT_SNAPSHOT_DIR = REPO_ROOT / "docs" / "report_snapshots"

REPORT_TITLE = "密捷：面向动态词元剪枝的双向隐私安全推理系统"
REPORT_EMAIL = "待参赛团队补充"
REPORT_DATE = "2026年5月20日"

FINANCE_COMPRESSION = "68.39%"

HARDWARE_BASELINE = {
    "cpu": "Intel Xeon Gold 6148 @ 2.40GHz，80 逻辑核（2 路，每路 20 核，超线程开启）",
    "memory": "62 GiB",
    "os": "Ubuntu 20.04.6 LTS",
    "kernel": "Linux 5.15.0-139-generic",
    "python": "Python 3.9.25",
    "spu": "SecretFlow SPU 0.9.3.dev20241118（Apache License 2.0）",
    "runtime_mode": "两方 2PC colocated localhost 原型链路，通信量来自同配置 fresh rerun 的 SPU LinkDetails 计数器",
}

MATH_DISPLAY_CENTER_TAB = Inches(3.15)
MATH_DISPLAY_RIGHT_TAB = Inches(6.05)

ROBUSTNESS_CASE_LABELS = {
    "transfer_encoding_chunked_blocked": ("Chunked 编码绕过", "Transfer-Encoding: chunked，绕过 Content-Length"),
    "oversized_content_length_best_effort_413": ("超大 Content-Length", "声明体积超过 5 MiB，验证 413 + 强制阻断"),
    "duplicate_field_blocked": ("重复字段名", "重复注入 multipart 字段，诱导字段集歧义"),
    "boundary_fanout_blocked": ("Boundary 扇出", "构造过多分片，放大 MIME 解析树"),
    "nested_multipart_blocked": ("嵌套 multipart", "把 share part 伪装成 multipart/mixed"),
    "oversized_json_part_blocked": ("超长 JSON", "client_quality_summary 超过 JSON 字节门"),
    "unexpected_field_set_blocked": ("字段集漂移", "追加未声明字段，验证精确字段集约束"),
    "boundary_param_whitespace_blocked": ("Boundary 参数混淆", "boundary 带空白与引号变体"),
    "boundary_param_whitespace_rejected": ("Boundary 参数混淆", "boundary 带空白与引号变体"),
    "malformed_part_header_blocked": ("畸形 Part Header", "缺失合法头格式，验证 header parser"),
    "header_null_byte_blocked": ("Header 空字节注入", "在 part header 中注入 NUL / 控制字符"),
    "multipart_header_null_byte_blocked": ("Header 空字节注入", "在 part header 中注入 NUL / 控制字符"),
    "non_empty_epilogue_blocked": ("非空 Epilogue", "closing boundary 后继续追加垃圾字节"),
    "utf16_json_charset_blocked": ("异常 JSON 字符集", "UTF-16LE 编码结构化元数据，验证严格 UTF-8 约束"),
    "truncated_body_blocked": ("截断请求体", "提前断流，验证 streaming body reader"),
    "duplicate_nonce_concurrent": ("重复 Nonce 并发重放", "同 nonce 并发多发，验证幂等与重放防护"),
    "duplicate_payload_different_nonce": ("同载荷换 Nonce", "相同 share 指纹 + 新 nonce，验证 payload 去重"),
    "per_ip_inflight_limit": ("同 IP 并发占满", "单一来源地址并发冲击已通过前置校验的执行配额"),
    "ip_window_rate_limit": ("短窗限频", "同 IP 短时间空载请求探测限流门"),
}

ROBUSTNESS_EVIDENCE_CODES = {
    "transfer_encoding_chunked_blocked": "A1",
    "oversized_content_length_best_effort_413": "A2",
    "duplicate_field_blocked": "A3",
    "boundary_fanout_blocked": "A4",
    "nested_multipart_blocked": "A5",
    "oversized_json_part_blocked": "A6",
    "unexpected_field_set_blocked": "A7",
    "boundary_param_whitespace_rejected": "A8",
    "malformed_part_header_blocked": "A9",
    "multipart_header_null_byte_blocked": "A10",
    "non_empty_epilogue_blocked": "A11",
    "utf16_json_charset_blocked": "A12",
    "truncated_body_blocked": "A13",
    "duplicate_nonce_concurrent": "G1",
    "duplicate_payload_different_nonce": "G2",
    "per_ip_inflight_limit": "G3",
    "ip_window_rate_limit": "G4",
}

ROBUSTNESS_SOURCE_LABELS = {
    "protocol_fuzz_evidence.json": "协议异常输入验证记录",
    "control_plane_guard_evidence.json": "控制面守卫验证记录",
    "protocol_fuzz_transport_and_structure_cases.json": "协议异常输入分批记录 1",
    "protocol_fuzz_boundary_and_header_cases.json": "协议异常输入分批记录 2",
    "protocol_fuzz_truncated_body_case.json": "协议异常输入分批记录 3",
}

LAYER_LABELS = {
    "http_request_body_gate": "HTTP 请求体入口门",
    "content_length_header_gate": "Content-Length 头部上限门",
    "raw_multipart_precheck": "raw multipart 预检",
    "mime_tree_validation": "MIME 结构校验",
    "json_bytes_gate": "JSON 字节长度门",
    "strict_json_decoder": "严格 JSON 解码钩子",
    "content_type_boundary_parser": "Content-Type boundary 解析",
    "multipart_header_parser": "multipart header 解析",
    "strict_utf8_json_decoder": "strict UTF-8 JSON 解码",
    "json_numeric_hooks": "JSON 数值钩子",
    "streaming_body_reader": "流式 body 读取器",
    "replay_guard_nonce_cache": "Nonce 重放缓存门",
    "payload_fingerprint_guard": "payload 指纹去重门",
    "replay_guard_payload_cache": "载荷重放缓存门",
    "ip_inflight_limit": "IP inflight 配额门",
    "per_ip_inflight_guard": "单 IP inflight 守卫",
    "global_inflight_guard": "全局 inflight 守卫",
    "ip_sliding_window_guard": "IP 滑动窗口限频",
    "tcp_force_close_allowed": "TCP 强制阻断兜底",
    "exact_field_set_gate": "精确字段集门",
    "not_applicable": "不适用",
}


def resolve_template() -> Path:
    candidates = []
    if REPORT_SNAPSHOT_DIR.exists():
        snapshot_candidates = sorted(
            (
                path / OUTPUT.name
                for path in REPORT_SNAPSHOT_DIR.iterdir()
                if path.is_dir() and path.name != "autobackups" and (path / OUTPUT.name).is_file()
            ),
            reverse=True,
        )
        candidates.extend(snapshot_candidates)
    candidates.append(OUTPUT)
    for candidate in candidates:
        resolved = candidate.resolve()
        if resolved.is_file():
            return resolved
    searched = "\n".join(f"- {candidate}" for candidate in candidates)
    raise FileNotFoundError(
        "canonical report state not found. Create a snapshot first, or ensure the latest report exists at one of:\n"
        f"{searched}"
    )


def set_run_font(run, font_name="宋体", size_pt=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def configure_paragraph(paragraph, *, first_indent=True, align=None, line_spacing=1.5):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if line_spacing == 1.5 else None
    pf.first_line_indent = Pt(24) if first_indent else Pt(0)
    if align is not None:
        paragraph.alignment = align


def add_body(doc: Document, text: str, *, first_indent=True):
    paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(paragraph, first_indent=first_indent)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12, False)
    return paragraph


def m_seq(*items):
    return ("seq", *items)


def m_text(text: str, style: str | None = None):
    return ("text", text, style)


def m_var(text: str):
    return m_text(text)


def m_plain(text: str):
    return m_text(text, "p")


def m_bold(text: str):
    return m_text(text, "b")


def m_sub(base, sub):
    return ("sub", base, sub)


def m_sup(base, sup):
    return ("sup", base, sup)


def m_subsup(base, sub, sup):
    return ("subsup", base, sub, sup)


def m_frac(num, den):
    return ("frac", num, den)


def m_delim(beg: str, end: str, content):
    return ("delim", beg, end, content)


def m_nary(chr_value: str, sub, sup, content):
    return ("nary", chr_value, sub, sup, content)


def m_acc(chr_value: str, base):
    return ("acc", chr_value, base)


def m_join(items, sep):
    parts = []
    for index, item in enumerate(items):
        if index:
            parts.append(sep)
        parts.append(item)
    return m_seq(*parts)


def m_paren(*items):
    return m_delim("(", ")", m_seq(*items))


def m_bracket(*items):
    return m_delim("[", "]", m_seq(*items))


def m_abs(*items):
    return m_delim("|", "|", m_seq(*items))


def m_call(name_expr, *args):
    return m_seq(name_expr, m_paren(m_join(list(args), m_plain(","))))


def _math_run(text: str, style: str | None = None):
    run = OxmlElement("m:r")
    if style:
        r_pr = OxmlElement("m:rPr")
        sty = OxmlElement("m:sty")
        sty.set(qn("m:val"), style)
        r_pr.append(sty)
        run.append(r_pr)
    text_node = OxmlElement("m:t")
    if " " in text:
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    return run


def _append_math_expr(container, expr):
    kind = expr[0]
    if kind == "seq":
        for item in expr[1:]:
            _append_math_expr(container, item)
        return
    if kind == "text":
        container.append(_math_run(expr[1], expr[2] if len(expr) > 2 else None))
        return
    if kind == "sub":
        node = OxmlElement("m:sSub")
        base_node = OxmlElement("m:e")
        sub_node = OxmlElement("m:sub")
        _append_math_expr(base_node, expr[1])
        _append_math_expr(sub_node, expr[2])
        node.append(base_node)
        node.append(sub_node)
        container.append(node)
        return
    if kind == "sup":
        node = OxmlElement("m:sSup")
        base_node = OxmlElement("m:e")
        sup_node = OxmlElement("m:sup")
        _append_math_expr(base_node, expr[1])
        _append_math_expr(sup_node, expr[2])
        node.append(base_node)
        node.append(sup_node)
        container.append(node)
        return
    if kind == "subsup":
        node = OxmlElement("m:sSubSup")
        base_node = OxmlElement("m:e")
        sub_node = OxmlElement("m:sub")
        sup_node = OxmlElement("m:sup")
        _append_math_expr(base_node, expr[1])
        _append_math_expr(sub_node, expr[2])
        _append_math_expr(sup_node, expr[3])
        node.append(base_node)
        node.append(sub_node)
        node.append(sup_node)
        container.append(node)
        return
    if kind == "frac":
        node = OxmlElement("m:f")
        num_node = OxmlElement("m:num")
        den_node = OxmlElement("m:den")
        _append_math_expr(num_node, expr[1])
        _append_math_expr(den_node, expr[2])
        node.append(num_node)
        node.append(den_node)
        container.append(node)
        return
    if kind == "delim":
        node = OxmlElement("m:d")
        d_pr = OxmlElement("m:dPr")
        beg = OxmlElement("m:begChr")
        end = OxmlElement("m:endChr")
        beg.set(qn("m:val"), expr[1])
        end.set(qn("m:val"), expr[2])
        d_pr.append(beg)
        d_pr.append(end)
        content_node = OxmlElement("m:e")
        _append_math_expr(content_node, expr[3])
        node.append(d_pr)
        node.append(content_node)
        container.append(node)
        return
    if kind == "nary":
        node = OxmlElement("m:nary")
        nary_pr = OxmlElement("m:naryPr")
        chr_node = OxmlElement("m:chr")
        chr_node.set(qn("m:val"), expr[1])
        nary_pr.append(chr_node)
        sub_node = OxmlElement("m:sub")
        sup_node = OxmlElement("m:sup")
        content_node = OxmlElement("m:e")
        _append_math_expr(sub_node, expr[2])
        _append_math_expr(sup_node, expr[3])
        _append_math_expr(content_node, expr[4])
        node.append(nary_pr)
        node.append(sub_node)
        node.append(sup_node)
        node.append(content_node)
        container.append(node)
        return
    if kind == "acc":
        node = OxmlElement("m:acc")
        acc_pr = OxmlElement("m:accPr")
        chr_node = OxmlElement("m:chr")
        chr_node.set(qn("m:val"), expr[1])
        acc_pr.append(chr_node)
        base_node = OxmlElement("m:e")
        _append_math_expr(base_node, expr[2])
        node.append(acc_pr)
        node.append(base_node)
        container.append(node)
        return
    raise ValueError(f"unsupported math expression kind: {kind}")


def _append_omath(paragraph, expr):
    omath = OxmlElement("m:oMath")
    _append_math_expr(omath, expr)
    paragraph._element.append(omath)
    return omath


def _sup_l():
    return m_paren(m_var("l"))


def _chw_index():
    return m_join([m_var("c"), m_var("h"), m_var("w")], m_plain(","))


def _xy_args():
    return [m_var("x"), m_var("y")]


def _s_l():
    return m_sup(m_var("s"), _sup_l())


def _k_l():
    return m_sup(m_var("K"), _sup_l())


def _tau_l():
    return m_sup(m_var("τ"), _sup_l())


def _e_l():
    return m_sup(m_var("e"), _sup_l())


def _pi_l():
    return m_sup(m_var("π"), _sup_l())


def _m_i_l():
    return m_subsup(m_var("m"), m_var("i"), _sup_l())


def _e_i_l():
    return m_subsup(m_var("e"), m_var("i"), _sup_l())


def _h_i_l():
    return m_subsup(m_var("h"), m_var("i"), _sup_l())


def _tilde_h_i_l():
    return m_subsup(m_acc("̃", m_var("h")), m_var("i"), _sup_l())


def _s_i_l():
    return m_subsup(m_var("s"), m_var("i"), _sup_l())


def _p_chw():
    return m_sub(m_var("p"), _chw_index())


def _mu_c():
    return m_sub(m_var("μ"), m_var("c"))


def _sigma_c():
    return m_sub(m_var("σ"), m_var("c"))


def _x_chw():
    return m_sub(m_var("x"), _chw_index())


def _r_chw():
    return m_sub(m_var("r"), _chw_index())


def _share0_chw():
    return m_sub(m_plain("share0"), _chw_index())


def _share1_chw():
    return m_sub(m_plain("share1"), _chw_index())


def _q_k():
    return m_sub(m_var("q"), m_var("k"))


def _delta_k():
    return m_sub(m_var("Δ"), m_var("k"))


def _q_k_client():
    return m_subsup(m_var("q"), m_var("k"), m_paren(m_plain("client")))


def _q_k_server():
    return m_subsup(m_var("q"), m_var("k"), m_paren(m_plain("server")))


def _l_i():
    return m_sub(m_var("L"), m_var("i"))


def _rgb():
    return m_plain("RGB")


def _call_L(*args):
    return m_call(m_var("L"), *args)


def _call_lap(*args):
    return m_call(m_plain("lap"), *args)


def _math_registry():
    quad = m_plain(" ")
    comma_space = m_plain(", ")
    minus = m_plain("−")
    ge = m_plain("≥")
    le = m_plain("≤")
    eq = m_plain("=")
    plus = m_plain("+")

    formula_31 = m_seq(
        _tau_l(),
        eq,
        m_call(m_plain("TopKBoundary"), _s_l(), _k_l()),
        m_plain(","),
        quad,
        _m_i_l(),
        eq,
        m_plain("1"),
        m_bracket(
            m_paren(m_join([_s_i_l(), m_var("i")], m_plain(","))),
            ge,
            _tau_l(),
        ),
    )
    formula_32 = m_seq(_tilde_h_i_l(), eq, _m_i_l(), m_plain("·"), _h_i_l())
    formula_32a = m_seq(_e_i_l(), eq, _s_i_l(), minus, m_plain("ε"), m_plain("·"), m_var("i"))
    formula_32b = m_seq(
        _pi_l(),
        eq,
        m_call(m_plain("BitonicSortDesc"), _e_l()),
        m_plain(","),
        quad,
        _m_i_l(),
        eq,
        m_plain("1"),
        m_bracket(
            m_var("i"),
            m_plain("∈"),
            m_sub(m_sup(m_var("π"), _sup_l()), m_seq(m_plain("1:"), _k_l())),
        ),
    )
    formula_33 = m_seq(
        _x_chw(),
        eq,
        m_call(
            m_plain("clip"),
            m_frac(
                m_seq(
                    m_frac(_p_chw(), m_plain("255")),
                    minus,
                    _mu_c(),
                ),
                _sigma_c(),
            ),
            m_plain("−2"),
            m_plain("2"),
        ),
    )
    formula_34 = m_seq(
        _share0_chw(),
        eq,
        _r_chw(),
        m_plain(","),
        quad,
        _r_chw(),
        m_plain("∼"),
        m_seq(m_plain("U"), m_bracket(m_join([m_plain("−2"), m_plain("2")], m_plain(",")))),
        m_plain(","),
        quad,
        _share1_chw(),
        eq,
        _x_chw(),
        minus,
        _share0_chw(),
    )
    formula_35 = m_seq(
        m_sub(m_var("H"), m_plain("audit")),
        eq,
        m_call(
            m_plain("SHA256"),
            m_join(
                [
                    m_plain("v7"),
                    m_plain("nonce"),
                    m_call(m_var("H"), m_plain("src")),
                    m_call(m_var("H"), m_var("x")),
                    m_call(m_var("H"), m_plain("share0")),
                    m_call(m_var("H"), m_plain("share1")),
                ],
                m_plain("∥"),
            ),
        ),
    )
    formula_36 = m_seq(
        m_var("X"),
        eq,
        m_plain("share0"),
        plus,
        m_plain("share1"),
        m_plain(","),
        quad,
        _rgb(),
        eq,
        m_var("X"),
        m_plain("⊙"),
        m_var("σ"),
        plus,
        m_var("μ"),
    )
    formula_37 = m_seq(
        _l_i(),
        eq,
        m_plain("0.299"),
        m_sub(m_var("R"), m_var("i")),
        plus,
        m_plain("0.587"),
        m_sub(m_var("G"), m_var("i")),
        plus,
        m_plain("0.114"),
        m_sub(m_var("B"), m_var("i")),
        m_plain(","),
        quad,
        m_plain("OverExp"),
        eq,
        m_frac(m_plain("1"), m_var("N")),
        m_nary(
            "∑",
            m_seq(m_var("i"), eq, m_plain("1")),
            m_var("N"),
            m_seq(m_plain("1"), m_bracket(_l_i(), ge, m_plain("0.95"))),
        ),
    )
    formula_38 = m_seq(
        _call_lap(*_xy_args()),
        eq,
        minus,
        m_plain("4"),
        _call_L(*_xy_args()),
        plus,
        _call_L(m_seq(m_var("x"), minus, m_plain("1")), m_var("y")),
        plus,
        _call_L(m_seq(m_var("x"), plus, m_plain("1")), m_var("y")),
        plus,
        _call_L(m_var("x"), m_seq(m_var("y"), minus, m_plain("1"))),
        plus,
        _call_L(m_var("x"), m_seq(m_var("y"), plus, m_plain("1"))),
    )
    formula_39 = m_seq(
        _delta_k(),
        eq,
        m_abs(_q_k_client(), minus, _q_k_server()),
        m_plain(","),
        quad,
        _delta_k(),
        le,
        m_sup(m_plain("10"), m_plain("−4")),
    )

    return {
        r"F_{\mathrm{less}}": m_sub(m_var("F"), m_plain("less")),
        r"F_{\mathrm{mux}}": m_sub(m_var("F"), m_plain("mux")),
        r"\tau^{(l)}": _tau_l(),
        r"e_i^{(l)}": _e_i_l(),
        r"\pi^{(l)}": _pi_l(),
        r"s^{(l)}": _s_l(),
        r"K^{(l)}": _k_l(),
        r"m_i^{(l)}": _m_i_l(),
        r"h_i^{(l)}": _h_i_l(),
        r"\tilde{h}_i^{(l)}": _tilde_h_i_l(),
        r"p_{c,h,w}": _p_chw(),
        r"\mu_c": _mu_c(),
        r"\sigma_c": _sigma_c(),
        r"x_{c,h,w}": _x_chw(),
        r"\mathrm{share0}_{c,h,w}": _share0_chw(),
        r"\mathrm{share1}_{c,h,w}": _share1_chw(),
        r"q_k": _q_k(),
        r"\Delta_k": _delta_k(),
        r"e_i^{(l)},\ \pi^{(l)}": m_join([_e_i_l(), _pi_l()], comma_space),
        r"h_i^{(l)},\ \tilde{h}_i^{(l)}": m_join([_h_i_l(), _tilde_h_i_l()], comma_space),
        r"\mu_c,\ \sigma_c": m_join([_mu_c(), _sigma_c()], comma_space),
        r"\mathrm{share0}_{c,h,w},\ \mathrm{share1}_{c,h,w}": m_join([_share0_chw(), _share1_chw()], comma_space),
        r"q_k,\ \Delta_k": m_join([_q_k(), _delta_k()], comma_space),
        r"N,\ L(x,y),\ \mathrm{lap}(x,y)": m_join([m_var("N"), _call_L(*_xy_args()), _call_lap(*_xy_args())], comma_space),
        r"\tau^{(l)}=\mathrm{TopKBoundary}\!\left(s^{(l)},K^{(l)}\right),\quad m_i^{(l)}=\mathbf{1}\!\left[(s_i^{(l)},i)\geq\tau^{(l)}\right]": formula_31,
        r"\tilde{h}_i^{(l)}=m_i^{(l)}\cdot h_i^{(l)}": formula_32,
        r"e_i^{(l)}=s_i^{(l)}-\varepsilon\cdot i": formula_32a,
        r"\pi^{(l)}=\mathrm{BitonicSortDesc}\!\left(e^{(l)}\right),\quad m_i^{(l)}=\mathbf{1}\!\left[i\in\pi^{(l)}_{1:K^{(l)}}\right]": formula_32b,
        r"x_{c,h,w}=\mathrm{clip}\!\left(\frac{p_{c,h,w}/255-\mu_c}{\sigma_c},-2,2\right)": formula_33,
        r"\mathrm{share0}_{c,h,w}=r_{c,h,w},\quad r_{c,h,w}\sim\mathcal{U}[-2,2],\quad \mathrm{share1}_{c,h,w}=x_{c,h,w}-\mathrm{share0}_{c,h,w}": formula_34,
        r"H_{\mathrm{audit}}=\mathrm{SHA256}\!\left(\mathrm{v7}\parallel \mathrm{nonce}\parallel H(\mathrm{src})\parallel H(x)\parallel H(\mathrm{share0})\parallel H(\mathrm{share1})\right)": formula_35,
        r"X=\mathrm{share0}+\mathrm{share1},\quad RGB=X\odot \sigma+\mu": formula_36,
        r"L_i=0.299R_i+0.587G_i+0.114B_i,\quad \mathrm{OverExp}=\frac{1}{N}\sum_{i=1}^{N}\mathbf{1}[L_i\geq 0.95]": formula_37,
        r"\mathrm{lap}(x,y)=-4L(x,y)+L(x-1,y)+L(x+1,y)+L(x,y-1)+L(x,y+1)": formula_38,
        r"\Delta_k=\left|q_k^{(\mathrm{client})}-q_k^{(\mathrm{server})}\right|,\quad \Delta_k\leq 10^{-4}": formula_39,
    }


MATH_EXPR_REGISTRY = _math_registry()


def _resolve_math_expr(text: str):
    expr = MATH_EXPR_REGISTRY.get(text)
    if expr is None:
        raise KeyError(f"unregistered math expression: {text}")
    return expr


def add_formula(doc: Document, text: str, label: str | None = None):
    paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(paragraph, first_indent=False, line_spacing=1.0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        MATH_DISPLAY_CENTER_TAB,
        alignment=WD_TAB_ALIGNMENT.CENTER,
        leader=WD_TAB_LEADER.SPACES,
    )
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        MATH_DISPLAY_RIGHT_TAB,
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.SPACES,
    )
    paragraph.add_run("\t")
    _append_omath(paragraph, _resolve_math_expr(text))
    if label:
        run = paragraph.add_run(f"\t{label}")
        set_run_font(run, "Times New Roman", 10.5, False)
    return paragraph


def add_inline_math(paragraph, text: str, *, height_pt: float = 13.2):
    del height_pt
    return _append_omath(paragraph, _resolve_math_expr(text))


def add_mixed_body(doc: Document, segments, *, first_indent=True):
    paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(paragraph, first_indent=first_indent)
    for kind, content in segments:
        if kind == "text":
            run = paragraph.add_run(content)
            set_run_font(run, "宋体", 12, False)
        elif kind == "math":
            add_inline_math(paragraph, content)
        else:
            raise ValueError(f"unsupported segment kind: {kind}")
    return paragraph


def add_table_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(paragraph, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.2)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 10.5, False)
    return paragraph


def add_reference_body(doc: Document, text: str):
    paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(paragraph, first_indent=False, line_spacing=1.2)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 10.5, False)
    return paragraph


def add_heading(doc: Document, text: str):
    paragraph = doc.add_paragraph(style=doc.styles["Heading 1"])
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", 22, True)
    return paragraph


def add_subheading(doc: Document, text: str, *, size_pt=12):
    paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(paragraph, first_indent=False)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", size_pt, True)
    return paragraph


def add_table(doc: Document, headers, rows, *, font_size=10, header_font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        run = paragraph.add_run(str(header))
        set_run_font(run, "宋体", header_font_size, True)
    for ridx, row in enumerate(rows, start=1):
        for cidx, value in enumerate(row):
            cell = table.rows[ridx].cells[cidx]
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            run = paragraph.add_run(str(value))
            set_run_font(run, "宋体", font_size, False)
    return table


def add_symbol_table(doc: Document, headers, rows, *, font_size=9.2, header_font_size=9.2):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        run = paragraph.add_run(str(header))
        set_run_font(run, "宋体", header_font_size, True)
    for ridx, row in enumerate(rows, start=1):
        formula_latex, meaning, where = row
        # symbol cell
        cell = table.rows[ridx].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        add_inline_math(paragraph, formula_latex, height_pt=12.5)
        # meaning cell
        cell = table.rows[ridx].cells[1]
        paragraph = cell.paragraphs[0]
        paragraph.text = ""
        run = paragraph.add_run(str(meaning))
        set_run_font(run, "宋体", font_size, False)
        # location cell
        cell = table.rows[ridx].cells[2]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        run = paragraph.add_run(str(where))
        set_run_font(run, "宋体", font_size, False)
    return table


def add_figure(doc: Document, image_path: Path, caption: str, *, width_inches=6.2):
    if not image_path.exists():
        return None
    paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(paragraph, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    caption_paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(caption_paragraph, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.2)
    run = caption_paragraph.add_run(caption)
    set_run_font(run, "宋体", 10, False)
    return caption_paragraph


def add_code_block(doc: Document, title: str, source_ref: str, intro: str, code: str):
    add_subheading(doc, title, size_pt=12)
    add_body(doc, f"代码位置：{source_ref}", first_indent=False)
    add_body(doc, intro)
    add_body(doc, "以下仅保留与正文论证直接相关的关键代码节选，已省略无关分支、重复检查与通用异常处理。", first_indent=False)
    paragraph = doc.add_paragraph(style=doc.styles["Normal"])
    configure_paragraph(paragraph, first_indent=False, line_spacing=1.08)
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(code)
    set_run_font(run, "Consolas", 8.2, False)
    return paragraph


def backup_existing_output(output_path: Path) -> Path | None:
    if not output_path.exists():
        return None
    REPORT_SNAPSHOT_DIR.mkdir(parents=True, exist_ok=True)
    autobackup_dir = REPORT_SNAPSHOT_DIR / "autobackups"
    autobackup_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = f"{output_path.stem}.before_regen_{stamp}{output_path.suffix}"
    backup_path = autobackup_dir / backup_name
    shutil.copy2(output_path, backup_path)
    manifest = {
        "created_at": datetime.now().astimezone().isoformat(),
        "source_path": str(output_path),
        "backup_path": str(backup_path),
        "source_size_bytes": output_path.stat().st_size,
        "source_sha256": sha256(output_path.read_bytes()).hexdigest(),
        "generator": str(Path(__file__).resolve()),
    }
    manifest_path = backup_path.with_suffix(".json")
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return backup_path


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _paragraph_style_value(element) -> str:
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        return ""
    return p_style.get(qn("w:val"), "")


def clear_from_abstract(doc: Document):
    body = doc._element.body
    children = list(body.iterchildren())
    start_index = None
    for index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        if _element_text(child).strip() != "摘要":
            continue
        if _paragraph_style_value(child).upper().startswith("TOC"):
            continue
        start_index = index
        break
    if start_index is None:
        return
    removable_children = [child for child in children[start_index:] if child.tag != qn("w:sectPr")]
    for child in removable_children:
        body.remove(child)


def _replace_paragraph_text(paragraph, text: str, *, font_name="宋体", size_pt=12, bold=False, align=None, first_indent=False):
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)
    configure_paragraph(paragraph, first_indent=first_indent, align=align)
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size_pt, bold)
    return paragraph


def sanitize_front_matter(doc: Document):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("作品名称："):
            _replace_paragraph_text(paragraph, f"作品名称：{REPORT_TITLE}", first_indent=False)
        elif text.startswith("电子邮箱："):
            _replace_paragraph_text(paragraph, f"电子邮箱：{REPORT_EMAIL}", first_indent=False)
        elif text.startswith("提交日期："):
            _replace_paragraph_text(paragraph, f"提交日期：{REPORT_DATE}", first_indent=False)
        elif text == "填写说明" or text.startswith("1. 所有参赛项目必须") or text.startswith("2. 作品报告采用A4纸") or text.startswith("3. 作品报告中各项目说明文字") or text.startswith("4. 作品报告模板里已经列的内容") or text.startswith("5. 为保证网评的公平、公正"):
            _replace_paragraph_text(paragraph, "", first_indent=False)


def rewrite_toc(doc: Document):
    toc_entries = [
        "摘要",
        "第一章 作品概述",
        "第二章 系统设计",
        "第三章 系统实现",
        "第四章 测试方案与结果分析",
        "第五章 创新性与局限性",
        "第六章 复现与参赛声明",
        "参考文献",
        "附录A 关键代码实现",
    ]
    paragraphs = doc.paragraphs
    toc_index = None
    abstract_index = None
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        compact_text = text.replace(" ", "")
        if toc_index is None and compact_text == "目录":
            toc_index = index
        elif toc_index is not None and text == "摘要":
            abstract_index = index
            break
    if toc_index is None or abstract_index is None or abstract_index <= toc_index:
        return
    content_indices = list(range(toc_index + 1, abstract_index))
    for slot_index, paragraph_index in enumerate(content_indices):
        paragraph = paragraphs[paragraph_index]
        text = toc_entries[slot_index] if slot_index < len(toc_entries) else ""
        _replace_paragraph_text(
            paragraph,
            text,
            font_name="宋体",
            size_pt=12,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_indent=False,
        )


def read_lines(path: Path):
    return path.read_text(encoding="utf-8").splitlines()


def extract_window(path: Path, anchor: str, *, before=0, after=12):
    lines = read_lines(path)
    line_no = None
    for index, line in enumerate(lines, start=1):
        if anchor in line:
            line_no = index
            break
    if line_no is None:
        raise ValueError(f"anchor not found in {path}: {anchor}")
    start = max(1, line_no - before)
    end = min(len(lines), line_no + after)
    return line_no, "\n".join(lines[start - 1 : end]).rstrip()


def _latest_json(pattern: str) -> Path | None:
    matches = sorted(RESULT_DIR.glob(pattern))
    return matches[-1] if matches else None


def _primary_result_json(kind: str) -> Path | None:
    exact_map = {
        "protocol_fuzz": RESULT_DIR / "protocol_fuzz_evidence.json",
        "control_plane_guard": RESULT_DIR / "control_plane_guard_evidence.json",
    }
    exact = exact_map.get(kind)
    if exact and exact.exists():
        return exact
    pattern_map = {
        "protocol_fuzz": ["protocol_fuzz_batch_*.json", "web_demo_protocol_fuzz_*.json"],
        "control_plane_guard": ["control_plane_guard_*.json", "web_demo_guard_stress_*.json"],
    }
    for pattern in pattern_map.get(kind, []):
        matches = sorted(path for path in RESULT_DIR.glob(pattern) if "_manual" not in path.stem)
        if matches:
            return matches[-1]
    return None


def validate_prerequisites():
    required_files = [
        CALIBRATION_FIGURE,
        TOPOLOGY_FIGURE,
        SEQUENCE_FIGURE,
        ROBUSTNESS_FIGURE,
        BEST_DEMO_JSON,
        COMM_PROFILE_JSON,
        MEDICAL_DYNAMIC_SUMMARY_JSON,
        PRIMARY_PROTOCOL_FUZZ_JSON,
        PRIMARY_GUARD_STRESS_JSON,
        SPU_LICENSE,
        SPU_MODIFICATIONS,
        MODIFICATION_AUDIT_DOC,
        CHANGE_AUDIT_DOC,
        THIRD_PARTY,
        REPRODUCE_DOC,
    ]
    missing = [str(path) for path in required_files if not path.exists()]
    if missing:
        raise FileNotFoundError(
            "second-revision report generation aborted because required evidence is missing:\n"
            + "\n".join(f"- {item}" for item in missing)
        )

    notice_missing = []
    for path in VENDORED_NOTICE_FILES:
        if not path.exists():
            notice_missing.append(f"{path} (missing file)")
            continue
        header = "\n".join(path.read_text(encoding="utf-8").splitlines()[:24])
        if "TransShield local" not in header:
            notice_missing.append(str(path))
    if notice_missing:
        raise RuntimeError(
            "second-revision report generation aborted because Modification Notice mapping is incomplete:\n"
            + "\n".join(f"- {item}" for item in notice_missing)
        )


def load_comm_profile():
    return json.loads(COMM_PROFILE_JSON.read_text(encoding="utf-8"))


def load_best_demo():
    return json.loads(BEST_DEMO_JSON.read_text(encoding="utf-8"))


def load_medical_dynamic_summary():
    return json.loads(MEDICAL_DYNAMIC_SUMMARY_JSON.read_text(encoding="utf-8"))


def format_percent(value, *, digits=4):
    value = float(value)
    if value <= 1.0:
        value *= 100.0
    return f"{value:.{digits}f}%"


def format_auc(value, *, digits=4):
    if value is None:
        return "未统一留存"
    return f"{float(value):.{digits}f}"


def format_seconds(value, *, digits=2):
    return f"{float(value):.{digits}f} 秒/样本"


def _display_method_name(method: str) -> str:
    mapping = {
        "Transshield（ours）": "密捷（本作品正式动态主线）",
        "Transshield（static control）": "密捷（固定结构静态对照线）",
        "MPCViT [3]": "MPCViT [3]",
        "DeiT-S Static": "DeiT-S 静态模型",
        "Original DynamicViT": "原始 DynamicViT",
    }
    return mapping.get(method, method)


def _parse_bool_privacy(note: str, method: str) -> str:
    if "Transshield" in method:
        return "是"
    return "否"


def _parse_bool_dynamic(method: str) -> str:
    lowered = method.lower()
    if "dynamic" in lowered or "ours" in lowered:
        return "是"
    return "否"


def load_baseline_rows(best_demo: dict):
    rows = []
    for item in best_demo["external_comparison"]["additional_rows"]:
        rows.append(
            [
                _display_method_name(item["method"]),
                _parse_bool_privacy(item.get("note", ""), item["method"]),
                _parse_bool_dynamic(item["method"]),
                format_percent(item["threshold_accuracy"]),
                format_auc(item.get("auc")),
                item["note"],
            ]
        )
    return rows


def load_secure_benchmark_rows(best_demo: dict):
    rows = []
    for item in best_demo["standardized_secure_benchmark"].get("comparisons", []):
        group_name = {
            "same_shape_operator_proxy": "同形状算子代理基准",
            "architecture_proxy": "异构结构代理基准",
        }.get(item["comparison_group"], item["comparison_group"])
        rows.append(
            [
                group_name,
                f"{item['module_comm_ratio_left_over_right']:.3f}x",
                f"{item['time_ratio_left_over_right']:.3f}x",
                item["scope_note"],
            ]
        )
    return rows


def _layer_to_cn(value: str) -> str:
    return LAYER_LABELS.get(value or "", value or "未标注")


def _return_text(item: dict) -> str:
    if "status" in item:
        status = item.get("status")
        error_code = item.get("error_code") or "ok"
        return f"HTTP {status} / {error_code}"

    details = item.get("details") or {}
    if not isinstance(details, dict):
        return "见详情摘要"
    results = details.get("results")
    if isinstance(results, list) and results:
        counts = {}
        for result in results:
            key = result.get("error_code") or str(result.get("status"))
            counts[key] = counts.get(key, 0) + 1
        ordered = [f"{name}×{count}" for name, count in counts.items()]
        return "；".join(ordered)
    if "first" in details and "second" in details:
        first = details["first"].get("status")
        second_code = details["second"].get("error_code") or details["second"].get("status")
        return f"首发 {first}；复发 {second_code}"
    if "rate_limited_ip" in details:
        return f"rate_limited_ip×{details.get('rate_limited_ip', 0)}"
    return "见详情摘要"


def _source_label(source_path: Path) -> str:
    return ROBUSTNESS_SOURCE_LABELS.get(source_path.name, source_path.name)


def _evidence_label(source_path: Path, item_name: str) -> str:
    evidence_code = ROBUSTNESS_EVIDENCE_CODES.get(item_name)
    if evidence_code:
        return f"{_source_label(source_path)} / {evidence_code}"
    return f"{_source_label(source_path)} / {item_name}"


def _robustness_row(item: dict, source_path: Path) -> list[str]:
    label, payload_feature = ROBUSTNESS_CASE_LABELS.get(
        item.get("name", ""),
        (item.get("name", ""), str(item.get("details", ""))[:32]),
    )
    return [
        label,
        payload_feature,
        _layer_to_cn(item.get("interception_layer", "")),
        _layer_to_cn(item.get("fallback_layer", "")),
        _return_text(item),
        (item.get("system_state") or {}).get("summary", "n/a"),
        _evidence_label(source_path, item.get("name", "")),
    ]


def load_robustness_rows():
    rows = []
    fuzz_path = _primary_result_json("protocol_fuzz")
    if fuzz_path and fuzz_path.exists():
        fuzz_payload = json.loads(fuzz_path.read_text(encoding="utf-8"))
        for item in fuzz_payload.get("results", []):
            rows.append(_robustness_row(item, fuzz_path))

    guard_path = _primary_result_json("control_plane_guard")
    if guard_path and guard_path.exists():
        guard_payload = json.loads(guard_path.read_text(encoding="utf-8"))
        for item in guard_payload.get("checks", []):
            rows.append(_robustness_row(item, guard_path))
    return rows


def load_robustness_summary():
    groups = []
    total_cases = 0
    total_passed = 0
    total_stable = 0
    total_fd_socket_clean = 0
    sources = [
        ("协议层异常输入", _primary_result_json("protocol_fuzz"), "results"),
        ("控制面守卫", _primary_result_json("control_plane_guard"), "checks"),
    ]
    for label, path, key in sources:
        if not path or not path.exists():
            continue
        payload = json.loads(path.read_text(encoding="utf-8"))
        items = payload.get(key, [])
        case_count = len(items)
        passed_count = sum(1 for item in items if item.get("passed"))
        stable_count = sum(1 for item in items if (item.get("system_state") or {}).get("stable"))
        fd_socket_clean_count = 0
        for item in items:
            delta = ((item.get("system_state") or {}).get("delta") or {})
            if delta.get("fd_count", 0) == 0 and delta.get("socket_fd_count", 0) == 0:
                fd_socket_clean_count += 1
        total_cases += case_count
        total_passed += passed_count
        total_stable += stable_count
        total_fd_socket_clean += fd_socket_clean_count
        groups.append(
            [
                label,
                str(case_count),
                f"{passed_count} / {case_count}",
                f"{fd_socket_clean_count} / {case_count}",
                f"{stable_count} / {case_count}",
                _source_label(path),
            ]
        )
    if groups:
        groups.append(
            [
                "合计",
                str(total_cases),
                f"{total_passed} / {total_cases}",
                f"{total_fd_socket_clean} / {total_cases}",
                f"{total_stable} / {total_cases}",
                "汇总",
            ]
        )
    return groups


def build_report(doc: Document):
    comm = load_comm_profile()
    best_demo = load_best_demo()
    medical_dynamic = load_medical_dynamic_summary()
    robustness_rows = load_robustness_rows()
    robustness_summary_rows = load_robustness_summary()
    baseline_rows = load_baseline_rows(best_demo)
    secure_benchmark_rows = load_secure_benchmark_rows(best_demo)

    medical_acc = format_percent(medical_dynamic["best_threshold_accuracy"])
    medical_argmax = format_percent(medical_dynamic["argmax_accuracy"])
    medical_threshold = f"{float(medical_dynamic['best_threshold']):.10f}"
    medical_sec_per_sample = format_seconds(comm["medical"]["sec_per_sample"])
    finance_sec_per_sample = format_seconds(comm["finance"]["sec_per_sample"])

    additional_rows = {
        item["method"]: item for item in best_demo["external_comparison"]["additional_rows"]
    }
    transshield_dynamic = additional_rows["Transshield（ours）"]
    transshield_static = additional_rows["Transshield（static control）"]
    original_dynamic = additional_rows["Original DynamicViT"]
    deit_static = additional_rows["DeiT-S Static"]
    mpcvit = additional_rows["MPCViT [3]"]
    medical_auc = format_auc(transshield_dynamic["auc"])
    dynamic_vs_static_delta = transshield_dynamic["threshold_accuracy"] - transshield_static["threshold_accuracy"]
    dynamic_vs_original_delta = transshield_dynamic["threshold_accuracy"] - original_dynamic["threshold_accuracy"]
    gap_to_mpcvit = best_demo["external_comparison"]["gap_to_mpcvit_reference"]["threshold_accuracy_gap"]
    gap_to_deit = deit_static["threshold_accuracy"] - transshield_dynamic["threshold_accuracy"]
    robustness_case_count = len(robustness_rows)

    add_heading(doc, "摘要")
    add_body(
        doc,
        "本作品面向医院侧数据与人工智能公司侧模型参数同时敏感的协同推理场景，构建了一套可运行、可验证、可交付的双向隐私安全推理系统。针对传统单向隐私保护难以同时满足数据合规与模型知识产权保护的矛盾，本作品以动态视觉 Transformer（DynamicViT）为载体，对动态词元剪枝边界进行协议友好重写，将原始模型中的删除式决策改写为基于安全比较与安全选择的密态可执行语义，并在不引入额外可信第三方的前提下完成两方安全计算（two-party computation, 2PC）原型落地。"
    )
    add_body(
        doc,
        "区别于只验证密态推理链路可执行性的演示原型，本作品在浏览器前端与服务端之间补充了轻量控制面闭环：浏览器工作线程（Web Worker）负责本地数据质量评估、审计哈希链与小端序（little-endian）分片序列化；服务端在进入安全处理器单元（Secure Processing Unit, SPU）密态前向计算前，统一完成多部件报文预检、结构化元数据字节长度门、分片哈希校验、非有限数与次正规数阻断、服务端数据质量复核及审计落盘。该闭环使系统不再默认输入天然可信，而是将协议层、解析层、张量层与审计层组织为一条可验证的风险控制链。"
    )
    add_body(
        doc,
        f"本作品的正式交付场景为医疗影像隐私推理：在 524 张全量验证集样本上取得 {medical_acc} 的阈值精度，较同仓固定结构静态对照线提升 {dynamic_vs_static_delta:.4f} 个百分点；与未经安全化边界重写的原始 DynamicViT 相比，阈值精度提升 {dynamic_vs_original_delta:.4f} 个百分点。端到端双向隐私原型在 32 条部署验证样本上的平均时延为 {medical_sec_per_sample}，双向总通信量为 {comm['medical']['dual_total_gib']:.2f} GiB。金融任务仅作为边界压力验证，不作为第二正式主线；当前 8 条压力样本与明文参考逐样本一致，平均时延为 {finance_sec_per_sample}。此外，本作品围绕协议层异常输入、重放与并发探针已补充 {robustness_case_count} 类黑盒验证证据，用于支撑安全边界与可服务性说明。"
    )

    add_heading(doc, "第一章 作品概述")
    add_subheading(doc, "1.1 背景与目标")
    add_body(
        doc,
        "在医疗、金融等高敏感场景中，数据使用方通常不愿以明文形式交出原始数据，模型提供方也不愿直接暴露明文模型参数。单向隐私保护机制难以同时满足双方诉求：只强调输入侧保护会暴露模型知识产权，只强调模型侧保护则会引发数据合规风险。进一步地，若模型采用按样本动态变化的剪枝路径，则依赖明文删除、阈值比较和变长前向结构的决策链很难直接迁入安全计算图，这也是许多安全推理系统被迫退化为固定结构模型的重要原因。"
    )
    add_body(
        doc,
        "本作品的设计目标不是单纯提升胸片分类精度，而是在双向隐私前提下保留动态剪枝能力，并将其落实到可运行的工程系统中。为此，本作品选择动态词元剪枝视觉 Transformer 方法 DynamicViT 作为主要载体，围绕词元打分、阈值比较、保留决策与安全执行语义之间的映射关系，完成从模型结构、协议路径到前后端控制面的联动设计。"
    )
    add_subheading(doc, "1.2 应用场景与交付边界")
    add_table(
        doc,
        ["部分", "当前定位", "报告口径"],
        [
            ["医疗影像", "唯一正式交付场景", f"动态安全剪枝与双向隐私推理；524 张验证集阈值精度 {medical_acc}；32 条样本平均时延 {medical_sec_per_sample}"],
            ["金融欺诈图像化编码", "边界压力验证场景", f"用于验证极端分布输入下的动态路由稳定性；8/8 一致；平均时延 {finance_sec_per_sample}；参数保留 {FINANCE_COMPRESSION}"],
            ["前后端控制面", "系统闭环组成", "浏览器工作线程、本地数据质量评估、服务端权威快检、审计哈希链与审计落盘"],
        ],
    )
    add_subheading(doc, "1.3 核心贡献")
    add_table(
        doc,
        ["维度", "核心贡献", "直接作用"],
        [
            ["算法层", "将 DynamicViT 的删除式词元剪枝改写为安全比较与安全选择语义", "使动态剪枝不退化为固定结构安全推理"],
            ["系统层", "构建数据方与模型方双向隐私的 2PC 推理链路，并将动态剪枝预测器 PredictorLG 保留在安全执行域内", "同时保护输入数据、模型参数与动态决策状态"],
            ["工程层", "补齐浏览器工作线程、服务端权威快检、审计哈希链、协议异常输入验证与复现脚本", "使报告叙述与代码实现、测试证据保持一致"],
        ],
    )

    add_heading(doc, "第二章 系统设计")
    add_subheading(doc, "2.1 总体架构")
    add_body(
        doc,
        "系统由医院终端、本地浏览器工作线程、跨域传输通道、业务侧前置网关、服务端权威快检层和两方安全执行域共同组成。明文图像仅在浏览器工作线程中完成解码、裁剪、归一化、数据质量评估与分片生成，跨越信任边界后仅传输密态分片（share）、结构化元数据和审计摘要；服务端在进入 SPU 计算前完成协议与张量合法性检查，随后由两方安全执行链路联合输出最终分类结果。"
    )
    add_body(
        doc,
        "图 2-1 采用按信任边界分区的部署拓扑：左侧为医院终端与本地明文处理域，中间为广域网与隔离缓冲区（wide area network / demilitarized zone, WAN/DMZ）传输边界，右侧为业务侧前置网关、服务端权威快检层与参与方安全执行域。图中显式标注了明文终止点和密态分片生成起点，并以差异化图例区分本地明文处理与跨边界密态通信。"
    )
    add_figure(
        doc,
        TOPOLOGY_FIGURE,
        "图2-1 系统物理与逻辑部署图：绿色实线表示本地明文处理，红色虚线与锁图标表示跨越信任边界的密态分片与张量通信。",
        width_inches=6.4,
    )
    add_subheading(doc, "2.2 威胁模型与安全边界")
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["保护对象", "用户输入数据、模型参数、中间动态决策状态、审计与控制面证据"],
            ["参与方", "医院终端浏览器、浏览器工作线程、业务侧前置网关、服务端快检模块、两方 SPU 节点"],
            ["攻击者能力", "可监听网络、构造畸形多部件表单报文（multipart/form-data）、伪造结构化元数据、重放一次性随机标识（nonce）、提交异常密态分片、实施并发探测"],
            ["安全假设", "两方半诚实；不引入额外可信第三方；前端不作为可信根，关键裁决以服务端复算为准"],
            ["不覆盖范围", "恶意参与方串谋、浏览器被完全控制、进入长耗时 SPU 后的主动断连感知、生产级分布式 DoS、防输出反演攻击"],
            ["防护证据", "协议层模糊测试、重放与并发守卫测试、服务端张量快检与审计落盘证据"],
        ],
    )
    add_table(
        doc,
        ["边界项", "当前覆盖", "当前不覆盖或限制"],
        [
            ["输入明文保护", "明文图像只在浏览器工作线程内解码、裁剪与分片，服务端仅接收密态 share 与结构化元数据", "不覆盖浏览器被完全控制或本地终端失陷后的明文窃取"],
            ["模型参数保护", "模型参数不向数据使用方以明文暴露，仅在两方安全执行域内参与计算", "不覆盖恶意参与方串谋、模型抽取与生产级侧信道分析"],
            ["动态决策隐私", "PredictorLG、阈值判断与并列分数决断保留在安全执行域内", "不覆盖进入安全图后的主动任务取消、长连接断开感知与更强攻击模型"],
            ["协议与输入防护", "原始多部件表单预检、JSON（JavaScript Object Notation）字节门、share 哈希与张量合法性快检、重放与并发守卫已落地", "不覆盖生产级分布式 DoS、跨节点全局限流与外部 WAF/网关联动"],
            ["审计与复现", "nonce、请求载荷（payload）、质量摘要与审计事件落盘，报告与脚本可回溯", "不覆盖长期合规归档平台、外部 SIEM 集成与跨系统追责链"],
        ],
        font_size=9,
        header_font_size=9,
    )
    add_body(
        doc,
        "上述边界说明用于避免将原型级控制面夸大为生产级防御体系。更完整的未覆盖范围、输出侧风险与原型部署边界，统一放在第 5.5 节“当前局限性”中说明。"
    )
    add_subheading(doc, "2.3 软件流程")
    add_body(
        doc,
        "图 2-2 给出了完整的软件流转时序。主线程仅负责本地影像加载、预览和请求编排；浏览器工作线程独立完成图像解码、裁剪、数据质量评估指标提取、审计链构建与分片序列化；服务端首先执行原始报文体、边界参数、报文头与多部件结构校验，再对结构化元数据、密态分片与重构张量进行权威快检，全部通过后方可进入 SPU 密态前向计算。最终结果回传时同时包含分类结论、质量裁决、审计摘要与控制面开销，使展示链路与安全链路保持一致。"
    )
    add_figure(
        doc,
        SEQUENCE_FIGURE,
        "图2-2 端到端软件流转时序图：明文仅停留在浏览器工作线程域内，跨域仅传输密态分片与审计摘要。",
        width_inches=6.4,
    )
    add_subheading(doc, "2.4 DynamicViT 密态改写")
    add_table(
        doc,
        ["原始 DynamicViT 操作", "安全计算困难", "本作品改写方式", "代价", "验证方式"],
        [
            ["直接删除 token", "变长结构难以进入固定图安全执行，且会暴露路径", "改写为保留/置零的掩码化表达（keep/zero），由安全选择完成保留决策", "保留了部分冗余计算", "本地明文参考路径与 secure replay 一致性对齐"],
            ["明文阈值比较", "数据相关分支会泄露剪枝边界", "将比较逻辑改写为安全比较语义", "增加比较通信与控制流复杂度", "部署阈值回代与输出一致性验证"],
            ["数据相关 Top-K 选择", "动态索引与并列分数处理不稳定", "采用编码索引跟踪与排序网络构造保留掩码（keep-mask）", "排序代价上升", "排序结果与后续决策链功能验证"],
            ["外部明文决定 keep-mask", "会退化为半隐私运行", "将 PredictorLG、阈值判断和并列分数决断（tie）保留在安全执行域内", "图规模扩大、执行时间增长", "双向隐私运行链与前端控制面联调"],
        ],
        font_size=9,
        header_font_size=9,
    )
    add_subheading(doc, "2.5 硬件部署与软件基准")
    add_table(
        doc,
        ["项目", "配置"],
        [
            ["CPU", HARDWARE_BASELINE["cpu"]],
            ["物理内存", HARDWARE_BASELINE["memory"]],
            ["操作系统", HARDWARE_BASELINE["os"]],
            ["内核版本", HARDWARE_BASELINE["kernel"]],
            ["Python", HARDWARE_BASELINE["python"]],
            ["安全计算底座", HARDWARE_BASELINE["spu"]],
            ["原型运行模式", HARDWARE_BASELINE["runtime_mode"]],
        ],
    )

    add_heading(doc, "第三章 系统实现")
    add_subheading(doc, "3.1 模型改造与两方安全推理链路")
    add_body(
        doc,
        "模型实现遵循“明文训练、密态执行语义重写、双向隐私推理落地”三层结构。训练阶段保留 DynamicViT 的词元打分与样本级动态边界能力；部署阶段将删除式剪枝语义改写为固定图可执行的掩码化表达；推理阶段由 SPU 与 OpenBumbleBee 两方安全推理集成层承担密态执行，仅对外返回最终分类结果。该实现路径的目标不是复现一条完全等价的可变长明文执行图，而是在固定图安全执行约束下尽可能保留动态剪枝边界的判别价值。"
    )
    add_body(
        doc,
        "对应到动态词元剪枝的核心计算关系，正式主线不再直接删除词元，而是先确定当前 stage 的保留边界，再生成固定形状的 keep-mask，并据此完成后续安全选择。关键关系见式（3-1）至式（3-2b）："
    )
    add_formula(doc, r"\tau^{(l)}=\mathrm{TopKBoundary}\!\left(s^{(l)},K^{(l)}\right),\quad m_i^{(l)}=\mathbf{1}\!\left[(s_i^{(l)},i)\geq\tau^{(l)}\right]", "(3-1)")
    add_formula(doc, r"\tilde{h}_i^{(l)}=m_i^{(l)}\cdot h_i^{(l)}", "(3-2)")
    add_formula(doc, r"e_i^{(l)}=s_i^{(l)}-\varepsilon\cdot i", "(3-2a)")
    add_formula(doc, r"\pi^{(l)}=\mathrm{BitonicSortDesc}\!\left(e^{(l)}\right),\quad m_i^{(l)}=\mathbf{1}\!\left[i\in\pi^{(l)}_{1:K^{(l)}}\right]", "(3-2b)")
    add_mixed_body(
        doc,
        [
            ("text", "其中，"),
            ("math", r"s^{(l)}"),
            ("text", " 表示第 l 个剪枝 stage 的词元得分向量，"),
            ("math", r"K^{(l)}"),
            ("text", " 表示该 stage 允许保留的词元数，"),
            ("math", r"\tau^{(l)}"),
            ("text", " 表示由排序边界导出的保留阈值，"),
            ("math", r"e_i^{(l)}"),
            ("text", " 表示引入索引微扰后的编码键，"),
            ("math", r"\pi^{(l)}"),
            ("text", " 表示经编码键双调排序得到的索引排列，"),
            ("math", r"m_i^{(l)}"),
            ("text", " 表示第 i 个词元的保留掩码。报告正文中的安全选择原语 "),
            ("math", r"F_{\mathrm{mux}}"),
            ("text", " 与安全比较原语 "),
            ("math", r"F_{\mathrm{less}}"),
            ("text", " 正是围绕这组关系完成的协议友好改写。"),
        ],
    )
    add_body(
        doc,
        "式（3-2a）与式（3-2b）对应当前密态执行中的并列分数稳定化策略：先通过编码键将词元得分与索引绑定，再在固定比较网络中执行按索引跟踪的双调排序，最后直接从排序结果的前 K 个索引构造 keep-mask。这样做避免了“先排序求阈值、再对全体词元执行一次阈值比较”的二次密态比较开销，同时保证并列分数情况下的保留决策具有确定性。"
    )
    add_body(
        doc,
        "图 3-1 从实现结构角度补充展示当前项目框架：上层对应医院终端、业务网关与双节点密态执行域；中层分别对应控制面闭环、DynamicViT 安全化重写组件与两方半诚实安全计算协议；底层则是秘密分享、比较/选择与审计摘要等密码学基础原语。该图用于说明第 2 章“系统设计”与第 3 章“系统实现”的衔接关系。"
    )
    if CURRENT_ARCH_FIGURE.exists():
        add_figure(
            doc,
            CURRENT_ARCH_FIGURE,
            "图3-1 密捷项目当前框架总体架构图。",
            width_inches=6.35,
        )
    add_subheading(doc, "3.2 端到端执行流程与模块协同")
    add_body(
        doc,
        "结合图 3-1 的当前项目框架，当前端到端链路可以压缩为三个安全阶段。第一阶段是终端本地预处理：浏览器工作线程完成影像解码、质量评估、审计链构建与秘密分享，安全目标是使原始影像和明文像素止于终端侧。第二阶段是服务端边界拦截：前置网关与权威快检模块在进入 SPU 之前完成报文、结构化元数据、share 字节流与重构张量合法性校验，安全目标是使异常输入、重放请求和畸形张量止于密态执行入口之外。第三阶段是 SPU 密态执行：PredictorLG、编码键排序、keep-mask 构造与前向分类全部保留在两方安全执行域内，安全目标是避免动态决策链退化为外部明文回放。"
    )
    add_body(
        doc,
        "因此，本节不再重复图 2-2 中的逐步动作，而是强调各阶段的安全职责分工：前端负责明文边界与分片生成，服务端负责权威阻断与审计固化，SPU 路径负责完成双向隐私约束下的动态安全推理。"
    )
    add_subheading(doc, "3.3 前端控制面实现")
    add_body(
        doc,
        "前端实现采用主线程与单实例浏览器工作线程分离的结构。主线程只负责文件句柄保留、页面状态维护和请求编排；工作线程内部完成图像头部尺寸嗅探、异常尺寸阻断、解码与中心裁剪、数据质量指标提取、源图与 share 审计哈希计算、小端序 32 位浮点（Float32）字节流序列化与可转移数组缓冲区传递。该设计避免了大体量数组在主线程上的结构化克隆开销，同时将高开销计算与页面交互解耦。"
    )
    add_body(
        doc,
        "前端计算过程中的关键数值变换与审计关系见式（3-3）至式（3-5）："
    )
    add_formula(doc, r"x_{c,h,w}=\mathrm{clip}\!\left(\frac{p_{c,h,w}/255-\mu_c}{\sigma_c},-2,2\right)", "(3-3)")
    add_formula(doc, r"\mathrm{share0}_{c,h,w}=r_{c,h,w},\quad r_{c,h,w}\sim\mathcal{U}[-2,2],\quad \mathrm{share1}_{c,h,w}=x_{c,h,w}-\mathrm{share0}_{c,h,w}", "(3-4)")
    add_formula(doc, r"H_{\mathrm{audit}}=\mathrm{SHA256}\!\left(\mathrm{v7}\parallel \mathrm{nonce}\parallel H(\mathrm{src})\parallel H(x)\parallel H(\mathrm{share0})\parallel H(\mathrm{share1})\right)", "(3-5)")
    add_mixed_body(
        doc,
        [
            ("text", "其中，"),
            ("math", r"p_{c,h,w}"),
            ("text", " 表示裁剪后图像在通道 c、位置 (h,w) 的原始像素值，"),
            ("math", r"\mu_c"),
            ("text", " 与 "),
            ("math", r"\sigma_c"),
            ("text", " 分别对应 ImageNet 均值与标准差，"),
            ("math", r"x_{c,h,w}"),
            ("text", " 为送入模型的归一化张量值。为避免浏览器端异常数值写入底层缓冲区，前端会先将归一化结果裁剪到 [-2, 2]，再以 little-endian Float32 方式写入 share 字节流；审计链则把源图哈希、归一化张量哈希与两份 share 哈希绑定到同一个 nonce 上。"),
        ],
    )
    add_subheading(doc, "3.4 服务端权威快检与审计")
    add_body(
        doc,
        "服务端在进入 SPU 之前，依次执行原始报文体读取上限、boundary 参数解析、多部件表单（multipart）原始结构预检、精确字段集约束、单字段 JSON 字节长度门、严格 UTF-8 结构化元数据解析、share 哈希校验、张量重构前内存对齐复制、非有限数与次正规数阻断、服务端数据质量复核，以及 nonce、payload 与来源 IP 维度的重放与并发守卫。审计层同时记录哈希摘要、控制面开销与事件落盘结果，以保证前端展示、服务端阻断与事后核查使用的是同一条证据链。"
    )
    add_body(
        doc,
        "服务端快检并非只核对客户端声明，而是围绕重构张量、质量摘要与容差关系进行权威复算。关键关系见式（3-6）至式（3-9）："
    )
    add_formula(doc, r"X=\mathrm{share0}+\mathrm{share1},\quad RGB=X\odot \sigma+\mu", "(3-6)")
    add_formula(doc, r"L_i=0.299R_i+0.587G_i+0.114B_i,\quad \mathrm{OverExp}=\frac{1}{N}\sum_{i=1}^{N}\mathbf{1}[L_i\geq 0.95]", "(3-7)")
    add_formula(doc, r"\mathrm{lap}(x,y)=-4L(x,y)+L(x-1,y)+L(x+1,y)+L(x,y-1)+L(x,y+1)", "(3-8)")
    add_formula(doc, r"\Delta_k=\left|q_k^{(\mathrm{client})}-q_k^{(\mathrm{server})}\right|,\quad \Delta_k\leq 10^{-4}", "(3-9)")
    add_mixed_body(
        doc,
        [
            ("text", "其中，服务端会先对 share 执行内存对齐复制，再将绝对值小于 1e-30 的极小值刷新为零，以切断次正规数带来的微代码异常路径；随后检查 share 与重构张量是否存在非有限值，以及 share 幅值是否超过 1e3。对客户端与服务端分别得到的质量摘要 "),
            ("math", r"q_k"),
            ("text", "，只在基础连续数值出现显著偏离时才标记可疑漂移，而不会仅凭离散风险标签不一致就直接判定篡改。"),
        ],
    )
    add_table_caption(doc, "表3-1 第三章关键符号说明")
    add_symbol_table(
        doc,
        ["符号", "含义", "出现位置"],
        [
            [r"s^{(l)}", "第 l 个剪枝 stage 的词元得分向量", "式(3-1)"],
            [r"K^{(l)}", "第 l 个剪枝 stage 允许保留的词元数", "式(3-1)"],
            [r"\tau^{(l)}", "由安全 Top-K 边界导出的保留阈值", "式(3-1)"],
            [r"e_i^{(l)},\ \pi^{(l)}", "编码键与经双调排序得到的索引排列", "式(3-2a)、式(3-2b)"],
            [r"m_i^{(l)}", "第 i 个词元在第 l 个 stage 的保留掩码", "式(3-1)、式(3-2)、式(3-2b)"],
            [r"h_i^{(l)},\ \tilde{h}_i^{(l)}", "原始词元表征与掩码后的词元表征", "式(3-2)"],
            [r"p_{c,h,w}", "裁剪后图像在通道 c、位置 (h,w) 的原始像素值", "式(3-3)"],
            [r"\mu_c,\ \sigma_c", "通道 c 的归一化均值与标准差", "式(3-3)、式(3-6)"],
            [r"x_{c,h,w}", "归一化后的模型输入张量", "式(3-3)、式(3-4)"],
            [r"\mathrm{share0}_{c,h,w},\ \mathrm{share1}_{c,h,w}", "针对输入张量构造的两份秘密分享分片", "式(3-4)"],
            [r"q_k,\ \Delta_k", "第 k 个质量指标及其客户端/服务端绝对偏差", "式(3-9)"],
            [r"N,\ L(x,y),\ \mathrm{lap}(x,y)", "像素总数、亮度值与离散拉普拉斯响应", "式(3-7)、式(3-8)"],
        ],
    )
    add_subheading(doc, "3.5 关键交付物与证据落点")
    add_table(
        doc,
        ["类别", "仓内落点", "用途"],
        [
            ["前后端演示程序", "tools/transshield_chat_demo.py；web_demo/control_plane_worker.js", "承载浏览器工作线程、服务端快检与页面交互"],
            ["自动化验证脚本", "tools/web_demo_protocol_fuzz.py；tools/web_demo_guard_stress.py", "验证协议层异常输入、重放与并发守卫"],
            ["图件与报告生成脚本", "tools/generate_report_figures.py；tools/generate_competition_report.py", "重建正式交付报告与图件"],
            ["复现说明", "README_REPRODUCE.md", "提供环境约束、运行步骤与预期输出"],
            ["许可证与修改映射", "spu_vendored/LICENSE；spu_vendored/MODIFICATIONS.md；THIRD_PARTY.md", "支撑第三方许可与本地修改说明"],
        ],
        font_size=9.2,
        header_font_size=9.2,
    )

    add_heading(doc, "第四章 测试方案与结果分析")
    add_subheading(doc, "4.1 测试环境、数据与评价指标")
    add_table(
        doc,
        ["对象", "样本规模", "作用", "主要指标"],
        [
            ["医疗全量验证集", "524 张", "正式阈值搜索与精度评估", "argmax 精度、阈值精度、部署阈值"],
            ["医疗部署验证批次", "32 张", "完整双向隐私原型运行", "平均时延、双向总通信量、隐私边界"],
            ["金融边界压力样本", "8 条", "极端分布输入下的稳定性验证", "逐样本一致性、时延、通信量"],
            ["协议层与控制面黑盒用例", f"{robustness_case_count} 类", "异常输入、重放与并发探针验证", "首个拦截层、兜底层级、资源状态"],
        ],
        font_size=9.2,
        header_font_size=9.2,
    )
    add_body(
        doc,
        "医疗任务的正式效果指标以阈值精度为主，辅以 argmax 精度（按最大响应类别直接决策的精度）、AUC（受试者工作特征曲线下面积）、部署阈值位置与端到端运行代价；金融任务只承担边界压力验证职责，因此不将其包装为与医疗对称的第二应用主线。2026 年 5 月 20 日的服务器复核已为动态正式主线补齐同口径 AUC；与此同时，原始 DynamicViT 一行统一回收到当前仓内可复现的 5 epoch 基线口径，不再沿用缺少配套 checkpoint 的旧表述。"
    )
    add_subheading(doc, "4.2 基线对比")
    add_table(
        doc,
        ["方案", "双向隐私", "动态剪枝", "阈值精度", "AUC（曲线下面积）", "说明"],
        baseline_rows,
        font_size=9,
        header_font_size=9,
    )
    add_body(
        doc,
        f"从同仓与外部对比结果看，本作品正式动态主线相较固定结构静态对照线提升 {dynamic_vs_static_delta:.4f} 个百分点，相较当前仓内可复现的原始 DynamicViT 基线提升 {dynamic_vs_original_delta:.4f} 个百分点，说明协议友好重写并未使动态路径退化为不可部署状态。另一方面，与同数据集强明文参考 MPCViT（面向多方计算的视觉 Transformer）相比，本作品当前仍存在 {gap_to_mpcvit:.4f} 个百分点的阈值精度差距；与同架构静态上界 DeiT-S（数据高效图像 Transformer 小型模型）相比，差距为 {gap_to_deit:.4f} 个百分点。这些差距说明本作品的价值主要在于“双向隐私 + 动态剪枝安全执行”的系统能力，而不是在明文精度指标上全面超过所有强明文参考。"
    )
    add_body(
        doc,
        "需要说明的是，外部对比表中“原始 DynamicViT”一行现统一采用当前仓内与服务器均可复现的原始明文基线口径。此前遗留的“20 epoch / 80.73%”旧值未能在最终交付仓与服务器环境中追溯到配套权重文件，因此本次修订不再继续引用该旧数值。"
    )
    add_subheading(doc, "4.3 医疗主要交付场景结果")
    add_table(
        doc,
        ["指标", "结果", "说明"],
        [
            ["全量验证集 argmax 精度", medical_argmax, "反映动态路径在默认决策边界下的直接分类结果"],
            ["全量验证集阈值精度", medical_acc, "基于 524 张验证样本完成统一阈值搜索后的正式指标"],
            ["全量验证集 AUC", medical_auc, "2026 年 5 月 20 日服务器按同口径重跑补齐"],
            ["部署阈值", medical_threshold, "动态路径单独校准后的正式部署阈值"],
            ["32 条样本验证批次平均时延", medical_sec_per_sample, "当前主要交付场景的端到端双向隐私原型口径"],
            ["32 条样本验证批次双向通信量", f"{comm['medical']['dual_total_gib']:.2f} GiB", "来自同配置重测记录的双向通信总量"],
            ["隐私边界", "服务端不接收明文像素值，模型参数不以明文暴露", "对外仅返回最终分类结果"],
        ],
    )
    add_body(
        doc,
        "图 4-1 对比了医疗动态安全剪枝路径与 DenseNet121 明文卷积基线的概率分布及阈值位置。结合 argmax 精度与阈值精度可以看出，当前动态路径的主要问题并非样本排序能力完全丧失，而是部署边界相对固定结构参考发生了系统性偏移。因此，针对动态路径单独校准部署阈值，是当前交付方案中的必要步骤。"
    )
    add_figure(
        doc,
        CALIBRATION_FIGURE,
        "图4-1 医疗动态安全剪枝路径与 DenseNet121 明文基线的概率分布与阈值偏移对比。",
        width_inches=6.3,
    )
    add_body(
        doc,
        f"同口径服务器复核表明，动态正式主线在 full-val CPU runtime-pruning reference 口径下的 AUC 为 {medical_auc}。这说明当前动态路径并未失去样本排序能力，影响部署效果的关键问题仍然是决策边界相对固定结构参考发生系统性偏移，因此动态路径单独阈值校准仍是正式部署中的必要步骤。"
    )
    add_subheading(doc, "4.4 性能与通信量分析")
    add_table(
        doc,
        ["场景", "样本数", "总时长", "平均时延", "双向总通信量", "每样本通信量"],
        [
            [
                "医疗主要交付场景",
                str(comm["medical"]["sample_count"]),
                f"{comm['medical']['elapsed_sec']:.2f} 秒",
                medical_sec_per_sample,
                f"{comm['medical']['dual_total_gib']:.2f} GiB",
                f"{comm['medical']['per_sample_gib']:.2f} GiB",
            ],
            [
                "金融边界压力验证",
                str(comm["finance"]["sample_count"]),
                f"{comm['finance']['elapsed_sec']:.2f} 秒",
                finance_sec_per_sample,
                f"{comm['finance']['dual_total_gib']:.2f} GiB",
                f"{comm['finance']['per_sample_gib']:.2f} GiB",
            ],
        ],
    )
    if secure_benchmark_rows:
        add_table(
            doc,
            ["代理基准对比组", "通信比值（左/右）", "时间比值（左/右）", "适用范围说明"],
            secure_benchmark_rows,
            font_size=8.8,
            header_font_size=8.8,
        )
    add_body(
        doc,
        "性能证据分为两层：其一是上述端到端双向隐私运行结果，用于说明正式医疗主线与金融压力样本在当前原型中的总时延和总通信量；其二是统一 secure benchmark 下的代理基准对比，用于说明算子替换本身的开销趋势。在固定同一 DeiT-S 形状的代理基准对比中，密捷的安全友好算子将通信压缩到外部基准算子的 0.149x，时间压缩到 0.528x，说明算子替换本身具有正向收益。另一方面，跨模型结构的异构结构代理基准只能用于说明结构尺度差异会显著放大绝对代价，不能与医疗全量验证主链路混写。"
    )
    add_body(
        doc,
        "当前时延与通信量偏高，主要由三类因素共同造成。第一，正式主线保留了动态词元剪枝的密态比较、排序与掩码化决策链，相比固定结构安全推理会引入额外的协议轮次。第二，当前运行口径是同机本地双节点（colocated localhost）2PC 原型链路，虽然两方节点位于同一服务器内，但张量交换、线性层通信与 softmax 归一化指数函数相关算子仍构成主要代价来源，不能将 localhost 原型误解为“几乎无通信成本”。第三，当前 Web 演示后端以正确性与边界防护为首要目标，未针对生产级吞吐做异步化、流水化或任务队列改造。因此，本报告只保留已直接落盘的端到端结果与算子级代理证据，不额外给出未经观测的阶段占比、P50/P95 延迟或生产级并发吞吐承诺。"
    )
    add_subheading(doc, "4.5 金融边界压力验证结果")
    add_body(
        doc,
        "需要说明的是，金融域原始输入并非自然图像，而是信用卡欺诈检测任务中的 30 维 PCA 特征。为了复用 DeiT-S 的块嵌入（Patch Embedding）结构，本作品先将一维金融特征通过拟图像空间连续性编码映射为 224×224 灰度图，使输入保留可供视觉 Transformer 利用的局部平滑先验。仓内冻结 bundle 记录表明，早期的 patch-based 与 DCT 编码未形成稳定可用结果，当前边界压力验证对应的是 v3 image-like smooth/contrast encoding。"
    )
    add_table(
        doc,
        ["指标", "结果", "说明"],
        [
            ["8 条压力样本逐样本一致性", "8 / 8 样本差异为 0", "与明文参考逐样本对齐"],
            ["部署效率", finance_sec_per_sample, "仅作为边界压力验证口径"],
            ["双向总通信量", f"{comm['finance']['dual_total_gib']:.2f} GiB", "同配置重测口径"],
            ["参数规模保留比例", FINANCE_COMPRESSION, "体现低秩压缩收益"],
        ],
    )
    add_body(
        doc,
        "金融部分仅承担方法迁移性与极端输入稳定性验证职责，不作为与医疗并列的独立交付场景。其价值在于：在不改变安全运行链基本结构的前提下，验证控制面、动态路由与安全执行语义在另一类高敏感任务中的边界表现，从而证明本作品的工程设计并非只对单一数据分布有效。"
    )
    add_subheading(doc, "4.6 协议层异常输入与鲁棒性验证")
    add_body(
        doc,
        "鲁棒性验证采用穷举边缘分支矩阵，而不是以单一拦截率概括系统表现。协议层测试覆盖分块传输编码（Transfer-Encoding: chunked）、超大报文长度声明（Content-Length）、boundary 参数混淆、重复字段、畸形分片报文头、空字节注入、非空尾部附加数据、UTF-16 编码结构化元数据、超长结构化元数据以及截断请求体；控制面测试覆盖随机一次性标识（Nonce）并发重放、相同载荷更换 Nonce、同一来源地址并发占满与短窗限频。每个用例均记录首个拦截层、兜底层级、返回处置与资源状态回落摘要，以同时证明拦截有效性与系统可服务性。"
    )
    add_body(
        doc,
        "其中，ΔFD 表示文件描述符数量变化，ΔSock 表示套接字文件描述符数量变化，ΔRSS 表示常驻内存（KiB）变化，ΔThr 表示线程数变化；四项指标均按“触发后观测值减去基线观测值”计算，用于判断异常输入处理后资源是否回落到稳态。"
    )
    if robustness_summary_rows:
        add_table(
            doc,
            ["测试组", "用例数", "按预期通过", "FD/Socket 无泄漏", "稳态回落用例", "证据文件"],
            robustness_summary_rows,
            font_size=8.9,
            header_font_size=9,
        )
        add_body(
            doc,
            "从当前留存结果看，协议层异常输入与控制面守卫共覆盖 17 类黑盒用例，17/17 均返回预期拦截或限制结果；全部用例在采样窗口内均未观察到文件描述符或套接字文件描述符的净增长。与此同时，部分并发与大负载探针会带来瞬时常驻内存抬升，因此本作品将其如实记录为“资源成本存在但句柄未泄漏”，而不是将其包装为零成本防护。"
        )
    if robustness_rows:
        add_table(
            doc,
            ["异常场景", "攻击载荷特征", "首个拦截层级", "兜底层级", "返回/处置", "系统状态", "证据来源"],
            robustness_rows,
            font_size=8.6,
            header_font_size=9,
        )
    else:
        add_body(doc, "当前未发现已落盘的鲁棒性测试结果文件，请先运行协议层模糊测试脚本与并发守卫验证脚本。")
    add_figure(
        doc,
        ROBUSTNESS_FIGURE,
        "图4-2 协议层异常输入与控制面兜底验证矩阵：同时显示首个拦截层级与文件描述符、套接字及常驻内存回落摘要。",
        width_inches=6.4,
    )
    add_body(
        doc,
        "资源状态采用可观测代理指标描述，而不作无法直接证明的绝对化承诺。本文将“系统仍可服务”限定为：请求结束后 inflight 占用计数回落，文件描述符数量与套接字文件描述符数量回到基线附近，未观察到连接句柄持续泄漏；常驻内存在部分已通过前置校验的压力用例中出现瞬时抬升，但重复异常注入测试后未观察到单调持续增长，因此仅将其表述为原型阶段的瞬时资源成本，而不宣称所有指标完全归零。"
    )
    add_subheading(doc, "4.7 结果分析")
    add_body(
        doc,
        "综合上述结果可以得到三个结论。第一，本作品正式动态主线的主要收益并非体现在超越强明文参考，而是体现在双向隐私约束下保留了可部署的动态剪枝能力，并通过阈值校准把全量验证集精度恢复到 92.7481%。第二，算子级 proxy 对比说明 secure-friendly 算子本身具有明显的通信与时间收益，但结构尺度差异仍然是当前端到端时延偏高的关键原因。第三，协议层异常输入与控制面探针在当前实现下能够被前置阻断，且未观察到文件描述符或套接字句柄的持续泄漏，这为系统的原型级可服务性提供了必要的证据支撑。"
    )

    add_heading(doc, "第五章 创新性与局限性")
    add_subheading(doc, "5.1 算法层创新")
    add_body(
        doc,
        "现有问题在于：原始 DynamicViT 的词元删除、阈值比较与数据相关 Top-K 选择依赖明文控制流，直接迁入安全计算图会同时遭遇变长结构、路径泄露和并列分数不稳定三类困难，因此常见方案通常选择放弃动态剪枝，退回固定结构安全推理。"
    )
    add_mixed_body(
        doc,
        [
            ("text", "本作品的做法是把 pruning boundary 从“删除式表达”重写为“掩码化表达”，将词元保留决策改写为基于安全选择原语的 "),
            ("math", r"F_{\mathrm{mux}}"),
            ("text", " 表达，将阈值判断改写为基于安全比较原语的 "),
            ("math", r"F_{\mathrm{less}}"),
            ("text", " 表达，并结合编码键双调排序（Encoded-Key Bitonic Sort）处理安全 Top-K（前 K 选择）与并列分数决断问题。这样做的意义在于，不仅让动态剪枝能够进入固定图安全执行环境，而且保留了按样本变化的剪枝决策能力，而不是用一个固定结构模型替代原始方法。"),
        ],
    )
    add_body(
        doc,
        "这一层的技术难点在于：既要避免数据相关分支泄露，又要在固定图语义下近似保留原始动态模型的判别边界；同时，排序与并列分数处理必须在安全执行环境内稳定完成，否则最终决策链会因 tie 或索引不稳定而失去可验证性。"
    )
    add_body(
        doc,
        "实验证据表明，该改写并未使动态路径退化为不可部署状态：正式动态主线相较固定结构静态对照线的阈值精度提升 0.7634 个百分点，相较未经安全化边界重写的原始 DynamicViT 提升 12.0181 个百分点。该层内容由第 2.4 节密态改写表、第 4.2 节基线对比、第 4.3 节医疗主线结果与第 4.7 节结果分析共同支撑。"
    )
    add_subheading(doc, "5.2 系统层创新")
    add_body(
        doc,
        "现有问题在于，不少安全推理系统虽然使用了密态执行框架，但会在外部明文环境中预先决定剪枝路径，再把结果送入密态推理后端。这类做法虽然降低了工程难度，却使动态状态、路径信息甚至部分模型行为暴露在安全执行域之外，本质上仍属于半隐私运行。"
    )
    add_body(
        doc,
        "本作品的做法是将动态剪枝预测器 PredictorLG、阈值判断与并列分数决断保留在安全执行域内，并以 SPU 与 OpenBumbleBee 组成两方安全计算链路，对输入数据、模型参数和中间动态决策状态同时进行保护。这样，系统不再依赖“外部先算路径、内部只回放”的简化策略，而是在双向隐私前提下完成完整的动态决策链执行。"
    )
    add_body(
        doc,
        "这一层的技术难点在于：既要在两方半诚实模型下维持输入与参数的双向隐私边界，又要让动态路径的关键控制流保留在密态环境内执行；如果处理不当，系统很容易退化为只保护输入或只保护模型的单向隐私方案。"
    )
    add_body(
        doc,
        "实验证据方面，当前端到端原型在 32 条医疗部署验证样本上的双向总通信量为 84.47 GiB，且服务端不接收明文像素值、模型参数不以明文暴露，对外仅返回最终分类结果；金融边界压力样本 8/8 与明文参考逐样本一致。该层内容由第 2.1 节总体架构、第 2.2 节威胁模型与安全边界、第 3.1 节两方安全推理链路、第 4.3 节医疗结果与第 4.5 节金融边界验证共同支撑。"
    )
    add_subheading(doc, "5.3 工程层特色与可验证交付能力")
    add_body(
        doc,
        "现有问题在于，很多 Web 演示只展示模型输出结果，默认将输入视为天然可信，缺少协议层异常输入阻断、审计链、资源状态回落记录和可复现实验入口。对于信息安全类作品而言，这类展示很难支撑“系统已形成工程闭环”的判断。"
    )
    add_body(
        doc,
        "本作品的做法是把浏览器工作线程、本地数据质量评估、审计哈希链、服务端权威快检、协议层异常输入模糊测试以及重放与并发守卫一起纳入交付范围。这里的重点不是把工程实现拔高为核心算法创新，而是保证评审能够看到：系统不仅能给出分类结果，还能证明输入经过了哪些检查、异常请求在何处被阻断，以及资源状态是否回到稳态。"
    )
    add_body(
        doc,
        "这一层的技术难点在于：前端要在不阻塞主线程的前提下完成高开销预处理与审计摘要生成，服务端要在进入 SPU 前完成原始报文、JSON 元数据、share 字节流和张量层的多级快检，同时还要把异常输入、重放与并发探针的结果转化为可复核的证据文件。"
    )
    add_body(
        doc,
        "实验证据方面，本作品已覆盖 17 类协议层与控制面黑盒用例，17/17 按预期返回，且采样窗口内未观察到文件描述符或套接字文件描述符净增长；同时仓内提供 10–20 分钟最低可复现路径，用于复核完整计算链路。该层内容由第 3.3 节前端控制面、第 3.4 节服务端权威快检与审计、第 4.6 节鲁棒性矩阵、第 6.1 节复现路径和 `README_REPRODUCE.md` 共同支撑。"
    )
    add_subheading(doc, "5.4 与常见方案对比")
    add_table(
        doc,
        ["对比对象", "常见做法", "本作品差异"],
        [
            ["固定结构安全推理", "放弃动态剪枝，直接改为固定结构模型", "保留按样本变化的剪枝决策，并把边界重写为可密态执行语义"],
            ["明文剪枝 + 密态推理", "外部明文环境先决定路径，密态后端仅回放结果", "剪枝决策链保留在安全执行域内完成，避免半隐私运行"],
            ["单向隐私推理", "只保护输入或只保护模型参数", "同时保护输入数据、模型参数和动态中间状态"],
            ["普通 Web 演示", "只展示结果页面，不提供输入守卫与异常验证", "加入服务端快检、审计链、模糊测试与重放/并发守卫证据"],
        ],
        font_size=9.2,
        header_font_size=9.2,
    )
    add_body(
        doc,
        "这组对比说明，本作品的主要价值不在于简单地“再做一条安全推理链”，而在于同时解决了动态剪枝保留、双向隐私成立、动态决策链不外泄以及工程交付可验证四个问题。"
    )
    add_subheading(doc, "5.5 当前局限性")
    add_body(
        doc,
        "本作品当前仍有五点需要明确说明的边界。第一，Web 演示后端采用同步阻塞式 `ThreadingHTTPServer` 原型，进入长耗时 SPU 计算后的主动断连感知与任务取消能力尚未具备，因此现有并发守卫不能等价替代生产级拒绝服务防护。第二，当前仓内仅固化了端到端通信量与算子级代理基准证据，未单独落盘正式医疗主线的阶段级时延剖面，因此尚不能给出更细粒度的性能分解。第三，金融任务当前只覆盖 8 条边界压力样本，足以支撑稳定性验证，但不足以单独构成第二应用主线。第四，当前安全模型主要建立在两方半诚实假设上，不覆盖恶意参与方串谋、模型抽取和更强侧信道对抗场景。第五，当前系统尚未系统评估输出结果本身可能带来的反推风险，因此输出侧隐私仍需在后续工作中单独建模与验证。"
    )

    add_heading(doc, "第六章 复现与参赛声明")
    add_subheading(doc, "6.1 最低可复现路径（10–20 分钟）")
    add_body(
        doc,
        "在模型目录与基础 Python 环境已就位的前提下，评审可在 10–20 分钟内完成最小闭环验证：先启动 Web 演示并在浏览器端加载一张本地医疗样本，确认系统在本地完成预处理与分片后返回分类、质量保障结论、审计摘要与控制面耗时；再运行协议层异常输入脚本与至少一项控制面守卫检查，验证异常输入阻断与并发/重放防护。完整命令与故障排查说明见 `README_REPRODUCE.md`。"
    )
    add_table(
        doc,
        ["步骤", "预估耗时", "入口", "预期输出"],
        [
            ["启动 Web 演示并加载一张本地医疗样本", "3–8 分钟", "artifacts/server_inference_friendly_pack/run_web_demo.sh", "浏览器在本地完成预处理与分片后返回分类结果、质量保障结论、审计摘要与控制面耗时"],
            ["运行协议层异常输入验证", "2–4 分钟", "tools/web_demo_protocol_fuzz.py", "生成 `protocol_fuzz_evidence.json`（或自定义输出文件），验证 multipart/JSON/截断请求体阻断"],
            ["运行至少一项控制面守卫检查", "2–6 分钟", "tools/web_demo_guard_stress.py --checks duplicate_nonce 或 inflight", "验证重复 nonce、重复载荷或并发占满守卫"],
        ],
        font_size=9.2,
        header_font_size=9.2,
    )
    add_subheading(doc, "6.2 数据、模型与权重合规声明")
    add_body(
        doc,
        "本作品对数据、模型结构与权重文件分别采用分层合规声明。数据层面，报告与演示不包含真实、未经授权的患者隐私信息；医疗结果仅用于研究验证与竞赛展示，不作为临床诊断依据；金融压力样本仅承担方法边界验证职责。模型层面，公开仓库保留模型结构、运行脚本、图件与报告生成链；权重层面，仓库当前不默认随 Git 分发全部完整权重文件，仅保留必要的 bundle 元数据与加载入口，完整权重需在授权范围内单独提供。"
    )
    add_table(
        doc,
        ["对象", "当前做法", "合规边界"],
        [
            ["测试数据", "仅使用仓内声明范围内的验证数据与压力样本", "不对外分发未经授权的原始敏感数据"],
            ["模型结构代码", "随仓保留并可复现运行链", "遵循仓内第三方来源与修改映射说明"],
            ["模型权重", "默认不随 Git 完整分发，仅保留 bundle 元数据与加载入口", "需由维护方按授权边界单独提供"],
            ["第三方预训练/上游组件", "按原许可证与来源映射保留", "不对上游本体主张排他性知识产权"],
        ],
        font_size=9,
        header_font_size=9,
    )
    add_subheading(doc, "6.3 第三方许可、原创边界与技术公开范围")
    add_body(
        doc,
        "本作品的知识产权声明按“上游底座、基于 Apache 2.0 的本地适配、团队原创部分”三层边界撰写。报告只描述仓内已实际完成的物理映射，不对未存在的许可证分发项或未落地的文件头声明作超前表述。"
    )
    add_table(
        doc,
        ["物理映射项", "仓内落点", "当前事实"],
        [
            ["上游许可证文本", "spu_vendored/LICENSE", "已保留 Apache License 2.0 文本"],
            ["vendored 变更总说明", "spu_vendored/MODIFICATIONS.md", "已列出当前交付版明确引用的本地适配文件"],
            ["修改文件头声明 1", "spu_vendored/libspu/spu.proto", "文件头含密捷项目本地修改声明"],
            ["修改文件头声明 2", "spu_vendored/libspu/mpc/cheetah/arithmetic.h", "文件头含密捷项目本地修改声明"],
            ["修改文件头声明 3", "spu_vendored/libspu/mpc/cheetah/arithmetic.cc", "文件头含密捷项目本地修改声明"],
            ["修改文件头声明 4", "spu_vendored/libspu/mpc/cheetah/protocol.cc", "文件头含密捷项目本地修改声明"],
            ["第三方来源总表", "THIRD_PARTY.md", "已保留 DynamicViT / OpenBumbleBee 来源与许可证映射"],
        ],
        font_size=9,
        header_font_size=9,
    )
    add_body(
        doc,
        "本作品的安全计算底座基于 SecretFlow SPU，当前已核对上游分发许可为 Apache License 2.0。当前核对到的上游分发物保留了许可证文件（LICENSE），未发现独立的 NOTICE 分发文件，因此交付物按事实保留许可证文本和本地修改说明，不虚构并不存在的附加分发项。本作品不对上游 SPU、OpenBumbleBee 与相关协议实现本体主张排他性知识产权。"
    )
    add_body(
        doc,
        "本作品的原创贡献主要集中在动态控制流映射策略、浏览器工作线程控制面、服务端权威快检、前后端审计链闭环、演示程序集成以及报告与图件生成链。当前交付物公开代码、图件、报告与运行脚本；涉及完整模型权重和数据集的部分，按授权边界与来源限制单独管理，不在公开仓库中默认再分发。"
    )

    add_heading(doc, "参考文献")
    references = [
        "[1] Dosovitskiy A, Beyer L, Kolesnikov A, et al. An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale[C]//International Conference on Learning Representations (ICLR). 2021.",
        "[2] Rao Y, et al. DynamicViT: Efficient Vision Transformers with Dynamic Token Sparsification[C]//Advances in Neural Information Processing Systems (NeurIPS). 2021.",
        "[3] Zeng W, et al. MPCViT: Searching for Accurate and Efficient MPC-Friendly Vision Transformer with Heterogeneous Attention[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision (ICCV). 2023.",
        "[4] Touvron H, Cord M, Douze M, Massa F, Sablayrolles A, Jégou H. Training data-efficient image transformers & distillation through attention[C]//Proceedings of the 38th International Conference on Machine Learning (ICML). 2021.",
        "[5] Huang G, Liu Z, van der Maaten L, Weinberger K Q. Densely Connected Convolutional Networks[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR). 2017.",
        "[6] Ma J, et al. SecretFlow-SPU: A Performant and User-Friendly Framework for Privacy-Preserving Machine Learning[C]//USENIX Annual Technical Conference (USENIX ATC). 2023.",
        "[7] Lu W, et al. BumbleBee: Secure Two-party Inference Framework for Large Transformers[C]//Network and Distributed System Security Symposium (NDSS). 2025.",
        "[8] Batcher K E. Sorting Networks and Their Applications[C]//Proceedings of the AFIPS Spring Joint Computer Conference. 1968: 307-314.",
        "[9] Yao A C. Protocols for Secure Computations[C]//23rd Annual Symposium on Foundations of Computer Science (FOCS). 1982: 160-164.",
        "[10] Masinter L. Returning Values from Forms: multipart/form-data[S]. RFC 7578. IETF, 2015.",
        "[11] National Institute of Standards and Technology. Secure Hash Standard (SHS)[S]. FIPS PUB 180-4. Gaithersburg, MD: NIST, 2015.",
        "[12] WHATWG. HTML Living Standard: Web Workers[EB/OL]. https://html.spec.whatwg.org/dev/workers.html, 2026-05-20.",
    ]
    for reference in references:
        add_reference_body(doc, reference)

    add_heading(doc, "附录A 关键代码实现")
    worker_path = REPO_ROOT / "web_demo" / "control_plane_worker.js"
    server_path = REPO_ROOT / "tools" / "transshield_chat_demo.py"
    fuzz_path = REPO_ROOT / "tools" / "web_demo_protocol_fuzz.py"
    add_body(
        doc,
        "本附录不再罗列完整源文件，而是按“公式—机制—证据”对应关系摘录最能代表系统实现要点的局部代码，用于说明浏览器端 little-endian share 序列化、SPU 内部安全排序与 keep-mask 构造、结构化元数据长度门以及黑盒 fuzz 用例构造等核心实现。",
        first_indent=False,
    )

    worker_line, _ = extract_window(worker_path, "function packFloat32LE", before=0, after=0)
    worker_code = textwrap.dedent(
        """\
        function packFloat32LE(values) {
          const bytes = new Uint8Array(values.length * 4);
          const view = new DataView(bytes.buffer);
          for (let index = 0; index < values.length; index += 1) {
            const value = values[index];
            if (!Number.isFinite(value)) throw new Error(`Non-finite float at ${index}`);
            view.setFloat32(index * 4, Math.fround(value), true);
          }
          return bytes;
        }"""
    )
    add_code_block(
        doc,
        "A.1 浏览器 Worker 中的 little-endian share 序列化",
        f"web_demo/control_plane_worker.js:{worker_line}",
        "该段代码展示了浏览器 Worker 如何显式使用 little-endian 方式写入 Float32 share，避免异构环境的字节序歧义。",
        worker_code,
    )

    spu_path = REPO_ROOT / "integrations" / "openbumblebee" / "e2e_secure_vit" / "spu_static_vit.py"
    secure_sort_line, _ = extract_window(spu_path, "def _bitonic_sort_desc_with_indices", before=0, after=0)
    secure_sort_code = textwrap.dedent(
        """\
        def _bitonic_sort_desc_with_indices(values):
            indices = jnp.broadcast_to(jnp.arange(N, dtype=jnp.int32), values.shape)
            ...
            left_index = jnp.where(is_left, p_arr, p_partner)
            is_desc = (left_index & k) == 0
            ...
            indices = jnp.where(has_partner, new_idx, indices)
            return x, indices

        def _secure_build_keep_decision(score, prev_decision_2d, keep_count):
            active_before = prev_decision_2d.squeeze(-1) > 0
            encoded_key = score - jnp.arange(N, dtype=score.dtype) * 1e-6
            encoded_masked = jnp.where(active_before, encoded_key, float('-inf'))
            ...
            sorted_keys, sorted_indices = _bitonic_sort_desc_with_indices(sortable)
            top_k_indices = sorted_indices[:, :keep_count]
            keep_mask = jnp.any(top_k_indices[:, :, None] == pos[None, None, :], axis=1)
            return (keep_mask[:, :N] & active_before)[:, :, None]"""
    )
    add_code_block(
        doc,
        "A.2 安全选择与双调排序的核心 SPU 实现",
        f"integrations/openbumblebee/e2e_secure_vit/spu_static_vit.py:{secure_sort_line}",
        "该段代码展示了 JAX/SPU 路径如何通过编码键、按索引跟踪的双调排序和向量化 keep-mask 构造，稳定处理安全 Top-K 与并列分数决断问题。",
        secure_sort_code,
    )

    json_line, _ = extract_window(server_path, "def _load_json_part", before=0, after=0)
    json_code = textwrap.dedent(
        """\
        def _load_json_part(self, raw_body: bytes, part: RawPart):
            part_view = memoryview(raw_body)[part.body_start:part.body_end]
            body_len = len(part_view)
            if body_len > JSON_PART_MAX_BYTES:
                raise ValueError('json part exceeds global size limit')
            field_limit = {...}.get(part.name, JSON_PART_MAX_BYTES)
            if body_len > field_limit:
                raise ValueError(f'{part.name} exceeds field size limit')
            decoded = bytes(part_view).decode('utf-8', 'strict')
            return json.loads(decoded, parse_int=strict_json_int,
                              parse_float=strict_json_float,
                              parse_constant=strict_json_constant)"""
    )
    add_code_block(
        doc,
        "A.3 JSON 字节长度门与 strict json.loads 钩子",
        f"tools/transshield_chat_demo.py:{json_line}",
        "该段代码展示了在 `json.loads` 前先执行字节长度门，并为整数、浮点和常量解析分别挂接严格钩子。",
        json_code,
    )

    fuzz_line, _ = extract_window(fuzz_path, "def case_header_null_byte", before=0, after=0)
    fuzz_code = textwrap.dedent(
        """\
        def case_header_null_byte(base_url: str, timeout: float, server_pid: int | None):
            boundary = f'nullbyte-{uuid.uuid4().hex}'
            header = (
                b'Content-Disposition: form-data; '
                b'name="domai\\x00n"'
            )
            body = build_custom_body(boundary, header, b'medical')
            raw = build_raw_multipart_request(
                base_url, boundary_param=f'boundary={boundary}', body=body
            )
            return capture_case(
                name='multipart_header_null_byte_blocked',
                action=lambda: send_raw_http(base_url, raw, timeout),
                interception_layer='multipart_header_parser',
                fallback_layer='mime_tree_validation',
            )"""
    )
    add_code_block(
        doc,
        "A.4 协议层异常输入黑盒验证样例",
        f"tools/web_demo_protocol_fuzz.py:{fuzz_line}",
        "该段代码展示了针对 multipart header 空字节注入的黑盒测试用例，以及脚本对拦截层级和资源状态的记录方式。",
        fuzz_code,
    )


def main():
    validate_prerequisites()
    template = resolve_template()
    doc = Document(str(template))
    sanitize_front_matter(doc)
    rewrite_toc(doc)
    clear_from_abstract(doc)
    build_report(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    backup_path = backup_existing_output(OUTPUT)
    doc.save(str(OUTPUT))
    if backup_path is not None:
        print(f"[backup] previous report saved to: {backup_path}")
    print(f"[ok] generated report: {OUTPUT}")


if __name__ == "__main__":
    main()
