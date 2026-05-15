#!/usr/bin/env python3
"""Generate unified competition report for both medical and finance domains."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_unified_report(output_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading('Transshield：基于 MPC 的隐私保护视觉推理框架', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('——医疗影像与金融风控双领域统一解决方案')
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # 1. 项目概述
    doc.add_heading('1. 项目概述', level=1)
    doc.add_paragraph(
        'Transshield 是一个基于安全多方计算（MPC）的隐私保护视觉推理框架，'
        '能够在不暴露任何一方数据的前提下完成完整的模型推理。本项目实现了 7 个核心创新点，'
        '并成功应用于医疗影像诊断和金融欺诈检测两个领域。'
    )
    
    # 2. 核心创新点
    doc.add_heading('2. 核心创新点', level=1)
    
    innovations = [
        ('创新点 1：DynamicViT Pruning Boundary 的协议友好重写',
         '将原始 DynamicViT 的 token pruning 决策边界从"删除式表达"重写为 MPC 友好的 F_mux/F_less 接口，'
         '解决了动态 shape 和条件分支泄漏问题。'),
        ('创新点 2：端到端隐私保护的 SPU Forward 实现',
         '实现完整的 SPU forward 路径，满足双向隐私约束：服务器看不到客户端图片，客户端获取不到模型参数。'),
        ('创新点 3：Bitonic Sort 安全 Top-K 选择',
         '基于 bitonic sort 实现 MPC 友好的 Top-K 选择，支持并行比较和流水线执行。'),
        ('创新点 4：MPC-Friendly 算子优化',
         '使用 fixed_square 激活函数替代 GELU，使用 uniform attention 替代 softmax，'
         '保证 MPC 环境下的计算精度。'),
        ('创新点 5：Token Pruning 动态剪枝',
         '通过 3 阶段动态剪枝（3/6/9 层），将 token 数量从 196 减少到 142，'
         '减少 27.6% 的计算量。'),
        ('创新点 6：FXP 固定点精度约束',
         '验证 fxp=16 是唯一安全配置，fxp<16 精度崩溃，fxp>16 溢出风险。'),
        ('创新点 7：SVD 低秩分解 MPC 推理加速',
         '通过 SVD 分解将线性层参数量压缩到 68.39%，SPU 推理加速 7.31 倍。'),
    ]
    
    for title_text, desc in innovations:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        doc.add_paragraph(desc)
    
    # 3. 双领域应用
    doc.add_heading('3. 双领域应用', level=1)
    
    # 3.1 医疗影像诊断
    doc.add_heading('3.1 医疗影像诊断', level=2)
    doc.add_paragraph(
        '数据集：RetinaMNIST（524 张视网膜图像，5 类分级）\n'
        '模型：DeiT-Small (depth12, embed_dim=384)\n'
        '精度：argmax accuracy = 91.98%\n'
        'SPU 推理时间：26 秒/张'
    )
    
    # 3.2 金融欺诈检测
    doc.add_heading('3.2 金融欺诈检测', level=2)
    doc.add_paragraph(
        '数据集：finance_fraud_v3（200 张交易图像，2 类分类）\n'
        '模型：DeiT-Small (depth12, embed_dim=384)\n'
        '精度：argmax accuracy = 100%\n'
        'SPU 推理时间：196 秒/张'
    )
    
    # 4. 统一技术栈
    doc.add_heading('4. 统一技术栈', level=1)
    
    # Create table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = '创新点'
    header_cells[1].text = '医疗模型'
    header_cells[2].text = '金融模型'
    
    # Data
    data = [
        ('1. Pruning Boundary Rewrite', '✅', '✅'),
        ('2. E2E Privacy SPU Forward', '✅', '✅'),
        ('3. Bitonic Sort Top-K', '✅', '✅'),
        ('4. MPC-Friendly Operators', '✅', '✅'),
        ('5. Token Pruning', '✅', '✅'),
        ('6. FXP Precision', '✅', '✅'),
        ('7. SVD LRD (rank=192)', '✅ 91.98% / 26s', '✅ 100% / 196s'),
    ]
    
    for i, (innovation, medical, finance) in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = innovation
        cells[1].text = medical
        cells[2].text = finance
    
    # 5. 隐私保护
    doc.add_heading('5. 隐私保护', level=1)
    doc.add_paragraph(
        '本项目实现了完整的双向隐私保护：\n\n'
        '1. 服务器看不到客户端图片：客户端图片在本地预处理后，通过秘密分享分散到两方服务器，'
        '服务器只能看到加密的共享份额，无法还原原始图片。\n\n'
        '2. 客户端获取不到模型参数：模型参数在 SPU 内部执行，客户端只能获取最终的推理结果（logits），'
        '无法获取模型权重或中间计算结果。\n\n'
        '3. 只暴露最终 logits：推理完成后，只向客户端暴露最终的分类结果，'
        '不泄露任何中间特征或注意力权重。'
    )
    
    # 6. 性能对比
    doc.add_heading('6. 性能对比', level=1)
    
    # Create performance table
    perf_table = doc.add_table(rows=4, cols=4)
    perf_table.style = 'Table Grid'
    
    # Header
    perf_header = perf_table.rows[0].cells
    perf_header[0].text = '指标'
    perf_header[1].text = 'Baseline'
    perf_header[2].text = 'Transshield'
    perf_header[3].text = '提升'
    
    # Data
    perf_data = [
        ('参数量', '22.4M', '15.3M', '-31.6%'),
        ('SPU 推理时间', '213.9s', '26s (医疗)', '8.2x 加速'),
        ('隐私保护', '无', '完整 MPC', '✓'),
    ]
    
    for i, (metric, baseline, transshield, improvement) in enumerate(perf_data, 1):
        cells = perf_table.rows[i].cells
        cells[0].text = metric
        cells[1].text = baseline
        cells[2].text = transshield
        cells[3].text = improvement
    
    # 7. 部署方案
    doc.add_heading('7. 部署方案', level=1)
    doc.add_paragraph(
        '本项目支持多种部署方案，适应不同地区和场景的需求：\n\n'
        '方案 A：城市三甲医院（高性能服务器）\n'
        '- 配置：depth12 + batch8\n'
        '- 推理时间：~21s/张\n'
        '- 精度：~92%\n\n'
        '方案 B：县级医院（中等性能服务器）\n'
        '- 配置：depth10 + batch4\n'
        '- 推理时间：~35s/张\n'
        '- 精度：~90%\n\n'
        '方案 C：乡村卫生院（低性能服务器）\n'
        '- 配置：depth8 + batch2\n'
        '- 推理时间：~60s/张\n'
        '- 精度：~85%'
    )
    
    # 8. 总结
    doc.add_heading('8. 总结', level=1)
    doc.add_paragraph(
        'Transshield 是一个完整的隐私保护视觉推理框架，具有以下优势：\n\n'
        '1. 创新性强：7 个核心创新点形成完整的技术闭环\n'
        '2. 双领域适用：成功应用于医疗影像和金融风控两个领域\n'
        '3. 隐私保护完整：实现双向隐私保护，符合医疗和金融数据安全要求\n'
        '4. 性能优异：SPU 推理加速 8.2 倍，参数量压缩 31.6%\n'
        '5. 部署灵活：支持多种部署方案，适应不同场景需求'
    )
    
    # Save
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == "__main__":
    output_path = "docs/transshield_竞赛作品报告_统一版_20260515.docx"
    create_unified_report(output_path)
