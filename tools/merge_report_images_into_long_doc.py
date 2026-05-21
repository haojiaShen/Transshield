#!/usr/bin/env python3
"""Safely merge generated figures into the long competition report docx.

This script only manipulates image paragraphs. It does not rewrite any正文文本.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


REPO_ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = REPO_ROOT / "docs" / "transshield_竞赛作品报告_最终版.docx"
BACKUP_PATH = REPO_ROOT / "docs" / "transshield_竞赛作品报告_最终版_图像合并前备份_20260520.docx"
ASSET_DIR = REPO_ROOT / "docs" / "report_evidence" / "assets"

FIGURES = [
    {
        "anchor_text": "| 隐私边界层 | 双向约束 | host_plaintext_pixel_values_materialized=false, host_model_params_materialized=false |",
        "image_path": ASSET_DIR / "system_trust_boundary_topology.png",
        "width": Inches(6.5),
    },
    {
        "anchor_text": "| 输出校准 | SPU-aware public logit-bias calibration |",
        "image_path": ASSET_DIR / "software_flow_sequence.png",
        "width": Inches(6.5),
    },
    {
        "anchor_text": "> 注：argmax accuracy 使用固定阈值 0.5（非校准阈值），实际部署使用校准阈值 0.358，对应 threshold accuracy 91.98%。",
        "image_path": ASSET_DIR / "medical_threshold_calibration_shift.png",
        "width": Inches(6.3),
    },
    {
        "anchor_text": "| 运行闭环 | ✅ | 批量稳定运行，失败恢复，结果落盘 |",
        "image_path": ASSET_DIR / "robustness_guard_matrix.png",
        "width": Inches(6.5),
    },
]


def _paragraphs_with_text(doc: Document):
    return list(doc.paragraphs)


def _remove_existing_drawing_paragraphs(doc: Document) -> None:
    body = doc._element.body
    for paragraph in list(doc.paragraphs):
        xml = paragraph._element.xml
        if "<w:drawing" in xml:
            body.remove(paragraph._element)


def _insert_image_after(paragraph, image_path: Path, width):
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_paragraph.add_run()
    run.add_picture(str(image_path), width=width)
    paragraph._element.addnext(new_paragraph._element)


def main():
    if not DOC_PATH.exists():
        raise FileNotFoundError(f"report docx not found: {DOC_PATH}")
    missing = [str(item["image_path"]) for item in FIGURES if not item["image_path"].exists()]
    if missing:
        raise FileNotFoundError(f"missing figure assets: {missing}")

    BACKUP_PATH.write_bytes(DOC_PATH.read_bytes())
    doc = Document(DOC_PATH)
    _remove_existing_drawing_paragraphs(doc)

    paragraphs = _paragraphs_with_text(doc)
    anchor_map = {}
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text:
            anchor_map[text] = paragraph

    for item in FIGURES:
        anchor_text = item["anchor_text"]
        paragraph = anchor_map.get(anchor_text)
        if paragraph is None:
            raise ValueError(f"anchor not found: {anchor_text}")
        _insert_image_after(paragraph, item["image_path"], item["width"])

    doc.save(DOC_PATH)
    print(f"[ok] merged report images into {DOC_PATH}")
    print(f"[ok] backup saved to {BACKUP_PATH}")


if __name__ == "__main__":
    main()
